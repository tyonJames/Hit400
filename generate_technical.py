"""
Generate BlockLand Technical Paper Word document.
Run: python generate_technical.py
Output: BlockLand_Technical_Paper.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY  = RGBColor(0x33, 0x33, 0x33)
LGRAY = RGBColor(0x66, 0x66, 0x66)
CODE_BG = '1E293B'

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '000000')
    pBdr.append(bot)
    pPr.append(pBdr)

def section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'{number}. {title.upper()}' if number else title.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK

def subsection_heading(doc, label, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{label} {title}')
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK

def body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK

def bullet(doc, text, sub=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.8 + (0.5 if sub else 0))
    run = p.add_run(f'{"o" if sub else "-"}  {text}')
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK

def code_block(doc, lines):
    """Render a code block as a shaded table with monospace font."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, 'F1F5F9')

    for li, line in enumerate(lines):
        if li == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph()

def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

def simple_table(doc, headers, rows, font_size=9):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], '000000')
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        bg = 'F5F5F5' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            set_cell_bg(cells[ci], bg)
            p = cells[ci].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(font_size)
    doc.add_paragraph()

# ── Paper sections ─────────────────────────────────────────────────────────────

def build_header(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BlockLand: Architecture and Design of a Dual-Ledger Land Registry System on the Stacks Blockchain')
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Tinotenda James')
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)

    for line in [
        'Department of Information Security and Assurance',
        'Harare Institute of Technology, Belvedere, Harare, Zimbabwe',
        'H220349M',
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = LGRAY
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()
    hr(doc)

def build_abstract(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    run = p.add_run('Abstract')
    run.bold = True
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    run = p.add_run(
        'This paper describes the architecture and key design decisions of BlockLand, a dual-ledger '
        'land registry system developed for Zimbabwe\'s land administration context. BlockLand '
        'combines a PostgreSQL relational database with the Stacks blockchain, exposed through a '
        'NestJS 10 REST API and a Next.js 14 frontend. The central design principle is a strict '
        'separation of concerns: the database stores operational metadata, while the blockchain '
        'stores only immutable ownership assertions enforced by a Clarity smart contract. The '
        'paper details the rationale for each architectural decision: the choice of Stacks over '
        'Ethereum, the Clarity language over Solidity, the dual-ledger approach over pure blockchain '
        'or pure database, and the specific design of the five on-chain data maps and three-step '
        'transfer workflow. Implementation details for the Clarity smart contract are presented '
        'with annotated code excerpts from the deployed contract. The design is evaluated against '
        'the functional requirements derived from Zimbabwe\'s Deeds Registries Act, demonstrating '
        'that each legal requirement maps to a verifiable on-chain constraint.'
    )
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('Keywords: ')
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(
        'blockchain, land registry, system architecture, Clarity, Stacks, NestJS, Next.js, '
        'PostgreSQL, dual-ledger, smart contract design, Zimbabwe'
    )
    run2.font.size = Pt(10)
    run2.italic = True

    hr(doc)

def build_introduction(doc):
    section_heading(doc, 'I', 'Introduction')

    body(doc,
        'BlockLand is a web-based land registry system designed to address the fraud, '
        'opacity, and administrative inefficiency of Zimbabwe\'s paper-based Deeds '
        'Registry. The system integrates a conventional web application stack with the '
        'Stacks blockchain as an immutable, publicly auditable ownership ledger. This '
        'paper focuses on the technical architecture and design decisions rather than '
        'security analysis (addressed in a companion paper) or literature context '
        '(addressed in the survey paper).',
        indent=True)

    body(doc,
        'The central design challenge was: how to obtain the immutability and '
        'transparency guarantees of a public blockchain while retaining the operational '
        'richness (complex queries, relationships, full-text search, file metadata) of '
        'a relational database — and to do so without sacrificing deployability in an '
        'environment with limited bandwidth and no gas-cost budget. Each architectural '
        'decision described in this paper is a response to a specific facet of this '
        'challenge.',
        indent=True)

def build_requirements(doc):
    section_heading(doc, 'II', 'Requirements and Design Principles')

    subsection_heading(doc, 'A.', 'Functional Requirements')
    body(doc,
        'Requirements were derived from two sources: Zimbabwe\'s Deeds Registries Act '
        '[Chapter 20:05], which specifies the legal obligations for property registration '
        'and transfer, and stakeholder interviews with intended user groups. The principal '
        'functional requirements are summarised in Table I.',
        indent=True)

    simple_table(doc,
        ['Req.', 'Description', 'Source', 'On-chain or Off-chain'],
        [
            ['FR-01', 'A registrar must approve every new property registration before it is committed to the official record', 'Deeds Registries Act s.13', 'On-chain: register-property requires registrar tx-sender'],
            ['FR-02', 'A property transfer requires consent from the seller, the buyer, and a registrar — independently', 'Deeds Registries Act s.16', 'On-chain: three-step smart contract workflow'],
            ['FR-03', 'No plot of land may be registered twice under different ownership records', 'Legal anti-fraud requirement', 'On-chain: title-deed-index prevents duplicate hash'],
            ['FR-04', 'Any citizen must be able to verify current ownership without a registered account', 'Public transparency requirement', 'Off-chain: /verify portal queries on-chain read-only functions'],
            ['FR-05', 'A complete, tamper-evident history of all past owners must be maintained', 'Audit and legal record requirement', 'On-chain: ownership-history map with composite key'],
            ['FR-06', 'A property under active dispute cannot be transferred', 'Legal protection requirement', 'On-chain: disputes map blocks initiate-transfer'],
            ['FR-07', 'Role-based access: only authorised registrars may register properties and finalise transfers', 'Deeds Registries Act s.4-5', 'On-chain: authorized-registrars map; Off-chain: RBAC middleware'],
        ]
    )
    caption(doc, 'TABLE I. Functional Requirements and Implementation Mapping')

    subsection_heading(doc, 'B.', 'Design Principles')
    for principle, explanation in [
        ('Blockchain as ground truth, database as cache.',
         'The blockchain is the authoritative source for ownership state. The PostgreSQL database '
         'stores a queryable copy of this state plus all operational metadata. If the two ever '
         'diverge, the blockchain record takes precedence. This one-directional consistency '
         'constraint is the most important invariant in the system.'),
        ('Minimal on-chain footprint.',
         'Only data that must be immutable and publicly verifiable is stored on-chain. '
         'Storing sale prices, notes, file paths, and timestamps off-chain reduces gas costs, '
         'protects commercially sensitive information, and avoids blockchain bloat.'),
        ('Decidable smart contract language.',
         'Clarity is chosen over Solidity specifically because it is non-Turing-complete: '
         'all execution paths terminate, reentrancy is architecturally impossible, and the '
         'contract can be fully analysed by static tools before deployment.'),
        ('Defence in depth.',
         'Access control is enforced at two independent layers — the NestJS API middleware '
         'and the Clarity contract. Both must be bypassed simultaneously for a privilege '
         'escalation to succeed.'),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(5)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f'{principle}  ')
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(explanation)
        run2.font.size = Pt(10)

def build_overview(doc):
    section_heading(doc, 'III', 'System Architecture Overview')

    body(doc,
        'BlockLand follows a three-tier architecture. Fig. 1 (described below in text '
        'form) illustrates the component relationships:',
        indent=True)

    simple_table(doc,
        ['Tier', 'Technology', 'Responsibility'],
        [
            ['Presentation', 'Next.js 14 (React, TypeScript, TailwindCSS)', 'User interface for all four roles; Hiro Wallet integration for blockchain tx signing'],
            ['Application', 'NestJS 10 (TypeScript, Fastify adapter)', 'REST API; JWT authentication; RBAC guards; business logic; blockchain transaction submission'],
            ['Persistence', 'PostgreSQL 15 + Stacks Blockchain (Clarity)', 'Operational metadata (DB); immutable ownership records (blockchain)'],
        ]
    )
    caption(doc, 'TABLE II. Three-Tier Architecture Summary')

    body(doc,
        'The NestJS backend is organised into domain modules following NestJS convention: '
        'AuthModule, UsersModule, PropertiesModule, TransfersModule, MarketplaceModule, '
        'DisputesModule, VerifyModule, and BlockchainModule. The BlockchainModule is the '
        'adapter layer between the application domain and the Stacks blockchain, '
        'encapsulating all use of the @stacks/transactions library.',
        indent=True)

    body(doc,
        'The frontend is organised into route groups: (auth) for login/register, '
        '(dashboard) for authenticated role-specific views, and public for the /verify '
        'portal accessible without authentication. Server components are used where '
        'possible (no client-side data fetching for static content); client components '
        'are used only where interactivity or wallet connection is required.',
        indent=True)

def build_data_layer(doc):
    section_heading(doc, 'IV', 'Data Layer Design')

    subsection_heading(doc, 'A.', 'PostgreSQL Schema')
    body(doc,
        'The PostgreSQL schema is managed via Prisma ORM. The principal tables and their '
        'relationships are:',
        indent=True)

    simple_table(doc,
        ['Table', 'Key Columns', 'Relationships'],
        [
            ['users', 'id (UUID), name, email, national_id, role (ENUM), wallet_address, is_active', 'owns many properties; party in many transfers'],
            ['properties', 'id (UUID), plot_number, address, area, property_type, owner_id (FK), blockchain_token_id (uint), ipfs_doc_hash, status (ENUM), registered_at', 'owned by one user; has many transfers; has many disputes'],
            ['transfers', 'id (UUID), property_id (FK), seller_id (FK), buyer_id (FK), registrar_id (FK), status (ENUM), blockchain_tx_hash, sale_price, initiated_at, completed_at', 'references property; three user FKs'],
            ['disputes', 'id (UUID), property_id (FK), raised_by (FK), description, evidence_ipfs_hash, status (ENUM), resolved_by (FK), resolution_notes', 'references property'],
            ['audit_log', 'id (UUID), actor_id (FK), action (VARCHAR), entity_type, entity_id, created_at', 'append-only; no FK enforcement on entity_id for flexibility'],
            ['marketplace_listings', 'id (UUID), property_id (FK), seller_id (FK), asking_price, status (ENUM), listed_at, removed_at', 'references property'],
            ['payments', 'id (UUID), transfer_id (FK), buyer_id (FK), amount, pop_file_path, verified_by (FK), created_at', 'references transfer; stores Proof of Payment path'],
        ]
    )
    caption(doc, 'TABLE III. PostgreSQL Schema Summary')

    subsection_heading(doc, 'B.', 'Dual-Ledger Consistency Protocol')
    body(doc,
        'Every state-changing operation that modifies ownership follows a strict '
        'two-phase protocol:',
        indent=True)

    for step, detail in [
        ('Phase 1 — Blockchain first.',
         'The NestJS service calls makeContractCall() from @stacks/transactions to '
         'construct the Clarity function call, signs it with the backend deployer wallet '
         'private key loaded from environment variables, and broadcasts it via '
         'broadcastTransaction(). The returned txid is stored immediately in the database '
         'as blockchain_tx_hash.'),
        ('Phase 2 — Database update on confirmation.',
         'The service polls getTransaction(txid) until the transaction status is '
         '"success". Only then is the PostgreSQL record updated to reflect the new state. '
         'If the blockchain transaction fails, the database is not updated, preserving '
         'consistency. The polling interval is 3 seconds with a 10-attempt timeout.'),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(5)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f'{step}  ')
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(detail)
        run2.font.size = Pt(10)

def build_contract(doc):
    section_heading(doc, 'V', 'Clarity Smart Contract Design')

    subsection_heading(doc, 'A.', 'Platform Choice: Stacks over Ethereum')
    body(doc,
        'Ethereum was evaluated as the initial blockchain platform but rejected for '
        'three reasons: (1) Solidity reentrancy and overflow vulnerabilities require '
        'significant defensive programming; (2) Ethereum gas costs during peak periods '
        'make land registry transactions economically prohibitive; (3) Ethereum finality '
        'depends on its own proof-of-stake consensus. Stacks was selected because its '
        'Clarity language is non-Turing-complete (eliminating reentrancy by design), '
        'transaction fees are denominated in STX which is significantly cheaper than '
        'ETH gas, and Stacks settles to Bitcoin, inheriting Bitcoin\'s 14-year '
        'track record of consensus security (Mansoor et al., 2023; Dwivedi et al., 2023).',
        indent=True)

    subsection_heading(doc, 'B.', 'On-Chain Data Maps')
    body(doc,
        'The contract (`blockland.clar`) defines seven data maps. The three most '
        'architecturally significant are shown below with their actual contract definitions:',
        indent=True)

    body(doc, 'Listing 1: property-registry map — the canonical ownership record.')
    code_block(doc, [
        '(define-map property-registry uint {',
        '  owner:          principal,',
        '  title-deed-hash: (buff 32),',
        '  status:         (string-ascii 20),',
        '  ipfs-doc-hash:  (optional (buff 32)),',
        '  registered-at:  uint',
        '})',
    ])

    body(doc,
        'The key is a uint (the PostgreSQL property ID, passed in at registration time). '
        'The value is a tuple storing the owner\'s Stacks principal, the SHA-256 hash of '
        'the physical title deed (as a 32-byte buffer), the property status, an optional '
        'IPFS document hash, and the Stacks block height at registration time. Block '
        'height serves as a blockchain-verifiable timestamp proxy: each Stacks block '
        'corresponds approximately to a 10-minute interval, providing a coarse but '
        'tamper-proof registration timestamp.',
        indent=True)

    body(doc, 'Listing 2: transfer-requests map — the three-step workflow state machine.')
    code_block(doc, [
        '(define-map transfer-requests uint {',
        '  seller:            principal,',
        '  buyer:             principal,',
        '  buyer-approved:    bool,',
        '  registrar-approved: bool',
        '})',
    ])

    body(doc,
        'At most one active transfer request may exist per property at any time '
        '(enforced by ERR-TRANSFER-PENDING). The state machine advances from '
        'buyer-approved: false (awaiting buyer) to buyer-approved: true (awaiting '
        'registrar). The entry is deleted on finalisation or cancellation — there '
        'is no "completed" status in the map because the completed state is simply '
        'the absence of a transfer-requests entry for that property.',
        indent=True)

    body(doc, 'Listing 3: title-deed-index — fraud prevention reverse lookup.')
    code_block(doc, [
        '(define-map title-deed-index (buff 32) uint)',
    ])

    body(doc,
        'This single-line map is the most important fraud prevention mechanism in the '
        'contract. Before registering a property, register-property checks: '
        '(is-none (map-get? title-deed-index title-deed-hash)). If a hash already '
        'exists in the index, the call fails with ERR-TITLE-DEED-EXISTS (u104). '
        'This makes it cryptographically impossible to register the same physical '
        'title deed document under two different property IDs.',
        indent=True)

    subsection_heading(doc, 'C.', 'Public Functions: register-property')
    body(doc,
        'The register-property function is the primary on-boarding entry point. '
        'It accepts four arguments — property-id (uint), owner (principal), '
        'title-deed-hash (buff 32), and ipfs-doc-hash (optional (buff 32)) — and '
        'enforces three ordered preconditions before writing any state:',
        indent=True)

    code_block(doc, [
        '(define-public (register-property',
        '  (property-id    uint)',
        '  (owner          principal)',
        '  (title-deed-hash (buff 32))',
        '  (ipfs-doc-hash  (optional (buff 32)))',
        ')',
        '  (begin',
        '    ;; CHECK 1: tx-sender must be an authorized registrar',
        '    (asserts! (check-is-registrar tx-sender) ERR-NOT-REGISTRAR)',
        '',
        '    ;; CHECK 2: property-id must not already exist',
        '    (asserts! (not (check-property-exists property-id)) ERR-PROPERTY-EXISTS)',
        '',
        '    ;; CHECK 3: title deed hash must not already be registered',
        '    (asserts! (is-none (map-get? title-deed-index title-deed-hash))',
        '              ERR-TITLE-DEED-EXISTS)',
        '',
        '    ;; Write to property-registry, title-deed-index, ownership-history',
        '    (map-set property-registry property-id {',
        '      owner: owner,  title-deed-hash: title-deed-hash,',
        '      status: "active",  ipfs-doc-hash: ipfs-doc-hash,',
        '      registered-at: block-height',
        '    })',
        '    (map-set title-deed-index title-deed-hash property-id)',
        '    (record-ownership-history property-id owner)',
        '    (var-set property-count (+ (var-get property-count) u1))',
        '    (ok property-id)',
        '  )',
        ')',
    ])

    body(doc,
        'The ordering of assertions is significant. Authorization (CHECK 1) is always '
        'verified before integrity checks (CHECKs 2-3). This ordering prevents an '
        'unauthorised caller from inferring information about existing property IDs or '
        'title deed hashes from error codes returned by integrity checks.',
        indent=True)

    subsection_heading(doc, 'D.', 'Transfer Workflow Functions')
    body(doc,
        'The three transfer functions form an atomic state machine. Listing 4 shows '
        'the initiate-transfer function, which creates the transfer-requests entry '
        'and locks the property into pending-transfer status:',
        indent=True)

    code_block(doc, [
        '(define-public (initiate-transfer (property-id uint) (buyer principal))',
        '  (let (',
        '    (property (unwrap! (map-get? property-registry property-id)',
        '                       ERR-PROPERTY-NOT-FOUND))',
        '  )',
        '    ;; Only the on-chain owner may initiate',
        '    (asserts! (is-eq tx-sender (get owner property)) ERR-NOT-OWNER)',
        '    ;; Buyer cannot be the same as seller',
        '    (asserts! (not (is-eq tx-sender buyer)) ERR-INVALID-BUYER)',
        '    ;; Property must not be disputed',
        '    (asserts! (not (default-to false (map-get? disputes property-id)))',
        '              ERR-PROPERTY-DISPUTED)',
        '    ;; No transfer may already be in progress',
        '    (asserts! (is-none (map-get? transfer-requests property-id))',
        '              ERR-TRANSFER-PENDING)',
        '',
        '    ;; Lock property and create transfer record',
        '    (map-set property-registry property-id',
        '      (merge property { status: "pending-transfer" }))',
        '    (map-set transfer-requests property-id {',
        '      seller: tx-sender, buyer: buyer,',
        '      buyer-approved: false, registrar-approved: false',
        '    })',
        '    (ok true)',
        '  )',
        ')',
    ])

    body(doc,
        'The buyer-approve-transfer function verifies tx-sender equals the designated '
        'buyer in the transfer record (not an argument — the contract reads the stored '
        'buyer principal) and sets buyer-approved to true. The '
        'registrar-finalize-transfer function verifies tx-sender is in '
        'authorized-registrars, verifies buyer-approved is true, atomically reassigns '
        'ownership in property-registry, appends an ownership-history entry, and '
        'deletes the transfer-requests entry. The ownership reassignment and '
        'history append occur within a single transaction — there is no intermediate '
        'state where ownership is ambiguous.',
        indent=True)

    subsection_heading(doc, 'E.', 'Read-Only Functions and Public Verification')
    body(doc,
        'Clarity\'s read-only functions change no state and require no transaction fee. '
        'The contract exposes five read-only functions used by the NestJS backend and '
        'the public /verify portal:',
        indent=True)

    simple_table(doc,
        ['Function', 'Parameters', 'Returns', 'Use'],
        [
            ['get-property-info', 'property-id (uint)', '(ok {owner, title-deed-hash, status, ...})', 'Property detail page; admin dashboard'],
            ['verify-owner', 'property-id (uint), claimed-owner (principal)', '(ok bool)', 'Public /verify portal ownership check'],
            ['get-transfer-request', 'property-id (uint)', '(optional {seller, buyer, buyer-approved, ...})', 'Transfer status polling'],
            ['get-ownership-history-entry', 'property-id (uint), seq (uint)', '(optional {owner, acquired-at})', 'Ownership history display'],
            ['is-registrar', 'address (principal)', '(ok bool)', 'Pre-validation before submitting registrar transactions'],
        ]
    )
    caption(doc, 'TABLE IV. Read-Only Contract Functions')

def build_api(doc):
    section_heading(doc, 'VI', 'API and Backend Design')

    subsection_heading(doc, 'A.', 'NestJS Module Architecture')
    body(doc,
        'The NestJS backend is organised into feature modules, each encapsulating a '
        'controller, service, and Prisma repository. The modules and their principal '
        'endpoints are:',
        indent=True)

    simple_table(doc,
        ['Module', 'Key Endpoints', 'Auth Required', 'Roles'],
        [
            ['AuthModule', 'POST /auth/register, POST /auth/login, POST /auth/refresh, POST /auth/logout', 'No / Yes', 'Public'],
            ['PropertiesModule', 'GET /properties, POST /properties, GET /properties/:id, GET /properties/:id/certificate', 'Yes', 'REGISTRAR (POST), OWNER/REGISTRAR (GET)'],
            ['TransfersModule', 'POST /transfers/initiate, POST /transfers/:id/approve, POST /transfers/:id/finalize, GET /transfers', 'Yes', 'OWNER (initiate), BUYER (approve), REGISTRAR (finalize)'],
            ['MarketplaceModule', 'GET /marketplace, POST /marketplace, DELETE /marketplace/:id', 'Partial', 'Public (GET), OWNER (POST/DELETE)'],
            ['DisputesModule', 'POST /disputes, PUT /disputes/:id/resolve', 'Yes', 'Any (POST), REGISTRAR (resolve)'],
            ['VerifyModule', 'GET /verify/:plotNumber, GET /verify/deed/:deedNumber, GET /verify/owner/:nationalId', 'No', 'Public'],
            ['UsersModule', 'GET /users, POST /users/:id/role', 'Yes', 'ADMIN'],
        ]
    )
    caption(doc, 'TABLE V. NestJS API Endpoints Summary')

    subsection_heading(doc, 'B.', 'Authentication Architecture')
    body(doc,
        'Authentication uses JWT RS256 asymmetric signing. The private key is loaded '
        'from an environment variable (PEM format). All API middleware verifies tokens '
        'using the corresponding public key. Two tokens are issued at login:',
        indent=True)

    for token, detail in [
        ('Access token (15-minute expiry).',
         'Transmitted in the Authorization: Bearer header. Stateless — the server '
         'holds no session state for access tokens. Verified by the JwtAuthGuard '
         'NestJS guard on every protected route.'),
        ('Refresh token (7-day expiry).',
         'Stored in the database with a revocation flag. Submitted to POST /auth/refresh '
         'to obtain a new access token. Revoked on logout and on anomaly detection.'),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f'{token}  ')
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(detail)
        run2.font.size = Pt(10)

    body(doc,
        'The JWT payload carries the user\'s id, email, and role. The RolesGuard '
        'decorator reads the role claim and compares it to the roles required by '
        'the route (specified via the @Roles() decorator). Roles are not looked '
        'up from the database on each request — the token is self-contained.',
        indent=True)

    subsection_heading(doc, 'C.', 'Blockchain Service Layer')
    body(doc,
        'The BlockchainService class encapsulates all interaction with the Stacks '
        'network. It exposes typed methods that map one-to-one to Clarity public '
        'functions: registerProperty(), initiateTransfer(), buyerApproveTransfer(), '
        'registrarFinalizeTransfer(), and isRegistrar(). Each method constructs '
        'a Clarity contract call using makeContractCall() from @stacks/transactions, '
        'signs it with the deployer private key, broadcasts it, and polls for '
        'confirmation before returning the txid. Read-only queries use '
        'callReadOnlyFunction() which returns results without broadcasting a transaction.',
        indent=True)

def build_frontend(doc):
    section_heading(doc, 'VII', 'Frontend Architecture')

    body(doc,
        'The Next.js 14 frontend uses the App Router. Route groups separate concerns: '
        '(auth) routes handle login and registration; (dashboard) routes are protected '
        'by a middleware that verifies the JWT access token and redirects to /login if '
        'missing or expired; the /verify route is fully public with no authentication '
        'check.',
        indent=True)

    simple_table(doc,
        ['Route Group', 'Principal Pages', 'Auth', 'Notes'],
        [
            ['(auth)', '/login, /register', 'No', 'Redirects to /dashboard if already authenticated'],
            ['(dashboard)/dashboard', 'Role-specific home page', 'Yes', 'Dashboard content conditionally rendered by role'],
            ['(dashboard)/properties', '/properties, /properties/[id], /properties/[id]/transfer', 'Yes', 'Owner and Registrar views differ'],
            ['(dashboard)/marketplace', '/marketplace, /marketplace/[id]', 'Partial', 'Listing requires auth; browsing does not'],
            ['(dashboard)/transfers', '/transfers, /transfers/[id]', 'Yes', 'Buyer sees pending approvals; Registrar sees finalisation queue'],
            ['(dashboard)/disputes', '/disputes, /disputes/[id]', 'Yes', 'Registrar resolution view'],
            ['public', '/verify', 'No', 'Queries blockchain read-only functions directly via API'],
        ]
    )
    caption(doc, 'TABLE VI. Next.js Route Structure')

    body(doc,
        'Wallet integration uses the @stacks/connect library. The ConnectWalletButton '
        'component triggers the Hiro Wallet browser extension (or Leather Wallet) to '
        'request connection. The connected wallet address is stored in client-side '
        'state and transmitted to the backend on registration to associate the user\'s '
        'Stacks principal with their BlockLand account. Future wallet-signing of '
        'transfer transactions will use the openContractCall() function from '
        '@stacks/connect, which presents the transaction to the user\'s wallet for '
        'signing rather than routing through the backend deployer key.',
        indent=True)

def build_evaluation(doc):
    section_heading(doc, 'VIII', 'Design Evaluation')

    body(doc,
        'The design is evaluated against the seven functional requirements in Table I. '
        'Table VII maps each requirement to its implementation and confirms fulfilment:',
        indent=True)

    simple_table(doc,
        ['Req.', 'Fulfilment Method', 'Verified By'],
        [
            ['FR-01', 'register-property asserts check-is-registrar(tx-sender); non-registrar calls return ERR-NOT-REGISTRAR', 'Clarinet unit test; adversarial test (Section V of companion security paper)'],
            ['FR-02', 'Three separate on-chain transactions required; each enforces tx-sender identity', 'End-to-end integration test; Clarinet test suite'],
            ['FR-03', 'title-deed-index reverse lookup; ERR-TITLE-DEED-EXISTS (u104) on duplicate', 'Clarinet unit test: duplicate-deed-rejected'],
            ['FR-04', 'verify-owner read-only function; /verify portal accessible without JWT', 'Manual browser test; no-auth API call returning on-chain state'],
            ['FR-05', 'ownership-history map with composite key {property-id, seq}; record-ownership-history called at registration and finalisation', 'Clarinet test: history-count increments correctly after each transfer'],
            ['FR-06', 'initiate-transfer checks (default-to false (map-get? disputes property-id)); returns ERR-PROPERTY-DISPUTED', 'Clarinet test: transfer-blocked-while-disputed'],
            ['FR-07', 'NestJS RolesGuard + @Roles() decorator (API layer); authorized-registrars map (contract layer)', 'API: 403 on role mismatch; Contract: ERR-NOT-REGISTRAR on non-registrar call'],
        ]
    )
    caption(doc, 'TABLE VII. Functional Requirements Fulfilment Matrix')

    body(doc,
        'One architectural limitation is noted: the current implementation routes all '
        'blockchain transactions through a backend deployer wallet rather than requiring '
        'individual users to sign from their own Hiro Wallet. This means the blockchain '
        'record carries the deployer\'s signature, not the individual user\'s. While the '
        'three-step workflow is enforced (the contract checks tx-sender for role validity), '
        'the non-repudiation property is weakened — a user could claim they did not '
        'personally authorise a transaction. Remediation (client-side wallet signing '
        'via openContractCall()) is identified as a priority for the next development '
        'iteration.',
        indent=True)

def build_conclusion(doc):
    section_heading(doc, 'IX', 'Conclusion')

    body(doc,
        'This paper has described the architecture and principal design decisions of '
        'BlockLand: a dual-ledger land registry system combining PostgreSQL with the '
        'Stacks blockchain via a NestJS API and Next.js frontend. The core architectural '
        'principle — blockchain as immutable ground truth, database as queryable cache — '
        'is operationalised through a strict two-phase consistency protocol and a Clarity '
        'smart contract that enforces all legally-mandated constraints directly at the '
        'blockchain layer.',
        indent=True)

    body(doc,
        'The choice of Clarity over Solidity eliminates reentrancy and unbounded-loop '
        'vulnerabilities by language design. The title-deed-index reverse lookup map '
        'provides a cryptographic guarantee against double-registration. The '
        'three-step transfer workflow enforces the seller-buyer-registrar consent '
        'model required by Zimbabwe\'s Deeds Registries Act without reliance on '
        'application-layer logic that could be bypassed.',
        indent=True)

    body(doc,
        'All seven principal functional requirements derived from the legal framework '
        'are fulfilled and verified through Clarinet unit tests and adversarial '
        'integration tests. The principal limitation — backend-mediated transaction '
        'signing — is acknowledged and targeted for resolution through client-side '
        'wallet signing in a future iteration.',
        indent=True)

def build_references(doc):
    section_heading(doc, '', 'References')
    hr(doc)

    refs = [
        'Alam, K.M., Rahman, J.M.A., Tasnim, A. and Akther, A. (2022) \'A Blockchain-based Land Title Management System for Bangladesh\', Journal of King Saud University - Computer and Information Sciences, 34(6), pp. 3096-3110.',
        'Dwivedi, R., Patel, S. and Shukla, S. (2023) \'Blockchain-Based Transferable Digital Rights of Land\', arXiv [Preprint], arXiv:2308.05950.',
        'Kimathi, K.T. (n.d.) A Land Administration Model Based on Blockchain Technology. [Unpublished thesis]. Harare: Harare Institute of Technology.',
        'Kombe, C., Manyilizu, M. and Mvuma, A. (2017) \'Design of Land Administration and Title Registration Model Based on Blockchain Technology\'.',
        'Konashevych, O. (n.d.) \'Constraints and Benefits of Blockchain Use for Real Estate and Property Rights\'. [Unpublished manuscript].',
        'M. Zein, R. and Twinomurinzi, H. (2023) \'Blockchain Technology in Lands Registration: A Systematic Literature Review\', JeDEM - eJournal of eDemocracy and Open Government, 15(2), pp. 1-36.',
        'Mansoor, M.A., Ali, M., Mateen, A., Kaleem, M. and Nazir, S. (2023) \'Blockchain Technology for Land Registry Management in Developing Countries\', 2023 2nd International Conference on ETECTE, Lahore, Pakistan.',
        'Paavo, J.P., Rodriguez-Puentes, R. and Chigbu, U.E. (2025) \'Practicality of Blockchain Technology for Land Registration: A Namibian Case Study\', Land, 14(8), p. 1626.',
        'Podshivalov, T.P. (2022) \'Improving implementation of the Blockchain technology in real estate registration\', Journal of High Technology Management Research, 33(2), p. 100440.',
        'Racetin, I., Kilic Pamukovic, J., Zrinjski, M. and Peko, M. (2022) \'Blockchain-Based Land Management for Sustainable Development\', Sustainability, 14(17), p. 10649.',
    ]

    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.8)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f'[{i}]  {ref}')
        run.font.size = Pt(9)

# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(10)

    build_header(doc)
    build_abstract(doc)
    build_introduction(doc)
    build_requirements(doc)
    build_overview(doc)
    build_data_layer(doc)
    build_contract(doc)
    build_api(doc)
    build_frontend(doc)
    build_evaluation(doc)
    build_conclusion(doc)
    build_references(doc)

    output = 'BlockLand_Technical_Paper.docx'
    doc.save(output)
    print(f'Saved: {output}')

if __name__ == '__main__':
    main()
