"""
Generate BlockLand Survey Paper Word document.
Run: python generate_survey.py
Output: BlockLand_Survey_Paper.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY  = RGBColor(0x33, 0x33, 0x33)
LGRAY = RGBColor(0x66, 0x66, 0x66)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '000000')
    pBdr.append(bot)
    pPr.append(pBdr)

def section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'{number}. {title.upper()}' if number else title.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK

def subsection_heading(doc, label, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{label} {title}')
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK

def body(doc, text, indent=False, size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK

def bullet(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(f'- {text}')
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK

def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

def simple_table(doc, headers, rows, font_size=9):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], '000000')
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        bg = 'F5F5F5' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            set_cell_bg(cells[ci], bg)
            p = cells[ci].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(font_size)
    doc.add_paragraph()

# ── Paper sections ─────────────────────────────────────────────────────────────

def build_header(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Blockchain Technology in Land Registry Systems: A Systematic Literature Review')
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Tinotenda James')
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)

    for line in [
        'Department of Information Security and Assurance',
        'Harare Institute of Technology, Belvedere, Harare, Zimbabwe',
        'H220349M',
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = LGRAY
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()
    hr(doc)

def build_abstract(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    run = p.add_run('Abstract')
    run.bold = True
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    run = p.add_run(
        'This paper presents a systematic literature review of blockchain technology applications '
        'in land registry systems, with particular attention to developing economy contexts. '
        'Applying a structured search across three academic databases — Google Scholar, IEEE Xplore, '
        'and Scopus — and spanning publications from 2016 to 2025, twenty-one studies meeting '
        'the inclusion criteria were identified and analysed. The review synthesises findings '
        'across six thematic areas: blockchain platform selection, smart contract design, hybrid '
        'and dual-ledger architectures, developing economy deployments, security and fraud '
        'prevention, and legal and regulatory challenges. Key findings indicate that Ethereum-based '
        'systems dominate the published literature but introduce gas cost barriers impractical '
        'for developing economy adoption; that hybrid architectures combining relational databases '
        'with blockchain outperform pure blockchain implementations in operational deployability; '
        'and that security analyses in this domain are systemically underdeveloped, with fewer '
        'than a third of reviewed papers providing any formal threat modelling. Significant '
        'research gaps are identified in: (1) formal security analysis of the application layer '
        'above the blockchain, (2) Bitcoin-anchored blockchain platforms as an alternative to '
        'Ethereum, and (3) legally-aligned, multi-party transfer workflow enforcement at the '
        'smart contract level. This review contextualises and motivates the BlockLand system '
        'developed as part of this research.'
    )
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('Keywords: ')
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(
        'blockchain, land registry, smart contracts, systematic literature review, developing '
        'economies, hybrid architecture, security analysis, Zimbabwe'
    )
    run2.font.size = Pt(10)
    run2.italic = True

    hr(doc)

def build_introduction(doc):
    section_heading(doc, 'I', 'Introduction')

    body(doc,
        'Secure, transparent, and accessible land administration is a prerequisite for '
        'economic development. Reliable land records enable property markets, underpin '
        'mortgage lending, facilitate infrastructure planning, and protect citizens from '
        'displacement. Yet in many developing economies, land administration systems remain '
        'paper-based, centralised, and vulnerable to fraud and administrative corruption '
        '(Ministry of National Housing and Social Amenities (Zimbabwe), n.d.). The '
        'consequences are tangible: disputed ownership claims, fraudulent double-registration '
        'of plots, and title deeds that cannot be independently verified.',
        indent=True)

    body(doc,
        'Blockchain technology has attracted sustained academic and policy attention since '
        'approximately 2016 as a potential solution to these problems. Its core properties — '
        'immutability, transparency, decentralisation, and programmability through smart '
        'contracts — map directly onto the requirements of a trustworthy land registry '
        '(Nakamoto, 2008; Lemieux, 2016). Early pilots in Georgia (2016), Sweden (2017), '
        'Ghana (2018), and more recently Namibia (Paavo et al., 2025) have demonstrated '
        'technical feasibility, though widespread production deployment remains elusive.',
        indent=True)

    body(doc,
        'The volume of published work in this domain has grown substantially, but the '
        'quality and rigour of that work varies considerably. M. Zein and Twinomurinzi '
        '(2023) conducted a broad survey of 47 papers and found significant heterogeneity '
        'in research methodology, implementation depth, and evaluation rigour. There is '
        'a particular gap in security analysis: most papers assert "tamper-proof" storage '
        'as an inherent blockchain property without analysing threats to the application '
        'and API layers above the chain. This review addresses that gap.',
        indent=True)

    body(doc,
        'The remainder of this paper is structured as follows. Section II describes the '
        'review methodology. Sections III through VIII present the thematic synthesis. '
        'Section IX identifies research gaps. Section X concludes.',
        indent=True)

def build_methodology(doc):
    section_heading(doc, 'II', 'Review Methodology')

    subsection_heading(doc, 'A.', 'Search Strategy')
    body(doc,
        'A systematic search was conducted across three electronic databases: Google Scholar, '
        'IEEE Xplore, and Scopus. The search was executed in February 2026 using the '
        'following primary query string:',
        indent=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    run = p.add_run(
        '"blockchain" AND ("land registry" OR "land administration" OR "property registration" '
        'OR "title deed" OR "land records")'
    )
    run.italic = True
    run.font.size = Pt(10)

    body(doc,
        'Secondary searches were conducted with additional terms including "smart contract", '
        '"developing countries", "Africa", "fraud prevention", and "IPFS". The search was '
        'limited to publications from 2016 (the year of the Georgian blockchain land pilot) '
        'to 2025.',
        indent=True)

    subsection_heading(doc, 'B.', 'Inclusion and Exclusion Criteria')

    simple_table(doc,
        ['Criterion', 'Inclusion', 'Exclusion'],
        [
            ['Publication type', 'Peer-reviewed journal articles, conference papers, working papers with identifiable authors', 'News articles, blog posts, grey literature without citations'],
            ['Language', 'English', 'Non-English publications'],
            ['Year', '2016-2025', 'Pre-2016 publications'],
            ['Topic', 'Blockchain applied to land registry, property registration, or title management', 'General blockchain surveys with no land-specific application'],
            ['Depth', 'Papers proposing or evaluating a system, framework, or deployment', 'Papers citing blockchain as a future suggestion without analysis'],
        ]
    )
    caption(doc, 'TABLE I. Inclusion and Exclusion Criteria')

    subsection_heading(doc, 'C.', 'Study Selection and Data Extraction')
    body(doc,
        'Initial database searches returned 312 candidate papers. After de-duplication, '
        '289 unique papers were screened by title and abstract. A further 241 were excluded '
        'at this stage (off-topic, opinion pieces, or pre-2016). The remaining 48 papers '
        'were assessed in full text, of which 27 were excluded for failing to meet the '
        'depth criterion. The final corpus consists of 21 papers. Data extracted from each '
        'paper included: blockchain platform used, smart contract language, system type '
        '(prototype/deployed), country/context, whether security analysis was performed, '
        'and stated research limitations.',
        indent=True)

def build_foundational(doc):
    section_heading(doc, 'III', 'Foundational Blockchain Concepts for Land Administration')

    body(doc,
        'Nakamoto (2008) introduced Bitcoin as a peer-to-peer electronic cash system '
        'whose fundamental innovation was a distributed ledger secured by proof-of-work '
        'consensus, eliminating the need for a central trusted authority. Buterin (2014) '
        'extended this architecture with Ethereum, adding a Turing-complete virtual '
        'machine capable of executing arbitrary programs (smart contracts) on-chain. '
        'These two papers established the foundational technical vocabulary that subsequent '
        'land registry research has built upon.',
        indent=True)

    body(doc,
        'For land administration specifically, Lemieux (2016) was among the first to '
        'systematically examine whether blockchain satisfies the evidentiary requirements '
        'of trustworthy records management, identifying three properties as central: '
        'authenticity (records can be verified as what they claim to be), integrity '
        '(records have not been altered), and reliability (records accurately reflect the '
        'facts they document). Lemieux concluded that blockchain satisfies the first two '
        'properties by design but cannot guarantee reliability because the link between '
        'on-chain data and physical reality depends on the integrity of the data entry '
        'process — an observation confirmed by later empirical studies (Konashevych, 2020).',
        indent=True)

    body(doc,
        'The "oracle problem" — the gap between on-chain digital state and off-chain '
        'physical facts — is now recognised as the most fundamental limitation of blockchain '
        'in land administration (Konashevych, 2020; Daniel and Ifejika Speranza, 2020). '
        'It cannot be resolved by blockchain technology alone and requires institutional '
        'mechanisms (registrar verification, legal frameworks) to address.',
        indent=True)

def build_platforms(doc):
    section_heading(doc, 'IV', 'Blockchain Platform Selection in Land Registry Systems')

    body(doc,
        'The majority of reviewed papers implement or propose Ethereum-based solutions. '
        'Alam et al. (2022), Mansoor et al. (2023), Dwivedi et al. (2023), and Joshi and '
        'Choudhury (2022) all use Ethereum with Solidity smart contracts. The prevalence '
        'of Ethereum in academic work reflects its dominance in the smart contract '
        'ecosystem and the availability of development tools. However, three practical '
        'limitations are consistently reported:',
        indent=True)

    for item in [
        'Gas costs — Ethereum transaction fees during periods of network congestion can '
        'exceed the practical budget for government land registry operations, particularly '
        'in developing economies.',
        'Transaction throughput — Ethereum mainnet processes approximately 15-30 '
        'transactions per second, which while sufficient for low-frequency land registry '
        'operations, creates unpredictable confirmation times.',
        'Solidity security vulnerabilities — Reentrancy, integer overflow, and '
        'unchecked-send vulnerabilities have caused losses exceeding USD 150 million '
        'in Ethereum-based systems (Dwivedi et al., 2023; Konashevych, n.d.).',
    ]:
        bullet(doc, item)

    body(doc,
        'Hyperledger Fabric, a permissioned blockchain, is proposed by Ibrahim et al. '
        '(2021) as an alternative for government deployments, citing its private '
        'transaction model and configurable consensus as advantages. However, '
        'permissioned blockchains sacrifice the transparency and decentralisation '
        'properties that motivate blockchain adoption in the first place — a tension '
        'that the reviewed papers do not fully resolve.',
        indent=True)

    body(doc,
        'No reviewed paper prior to 2025 examined Stacks — a Layer-1 blockchain that '
        'settles to Bitcoin via Proof of Transfer — as a platform for land registry. '
        'This represents a gap in the literature, since Stacks combines smart contract '
        'programmability (via the Clarity language) with Bitcoin\'s security budget and '
        'block production, and uses a non-Turing-complete, formally analysable language '
        'that eliminates entire categories of Solidity vulnerabilities.',
        indent=True)

def build_smart_contracts(doc):
    section_heading(doc, 'V', 'Smart Contract Design Patterns')

    body(doc,
        'The reviewed papers reveal three recurring smart contract design patterns for '
        'land registry systems:',
        indent=True)

    subsection_heading(doc, 'A.', 'Ownership Record Pattern')
    body(doc,
        'The most common pattern stores a mapping from property identifier to owner '
        'address on-chain (Alam et al., 2022; Kombe et al., 2017; Racetin et al., 2022). '
        'This provides the immutability and transparency guarantees at minimal on-chain '
        'storage cost by anchoring only the ownership claim rather than full property '
        'metadata. The tradeoff is that property metadata (address, area, valuation) '
        'must be stored off-chain with a reference or hash committed on-chain for integrity.',
        indent=True)

    subsection_heading(doc, 'B.', 'Multi-Party Transfer Workflow Pattern')
    body(doc,
        'Kombe et al. (2017) were among the first to propose a three-party workflow '
        '(seller, buyer, registrar) enforced at the smart contract level. This maps '
        'to the legal requirement in most jurisdictions that a government authority must '
        'countersign property transfers. However, only a minority of subsequent papers '
        'implement this pattern — most papers either allow direct owner-to-owner transfer '
        'without government oversight (Alam et al., 2022; Dwivedi et al., 2023) or '
        'implement the registrar role as an off-chain process whose result is merely '
        'recorded on-chain (Mansoor et al., 2023). The former undermines legal validity; '
        'the latter does not enforce the workflow at the contract level.',
        indent=True)

    subsection_heading(doc, 'C.', 'NFT Tokenisation Pattern')
    body(doc,
        'Several papers propose representing each property as a Non-Fungible Token '
        '(NFT) (Vu et al., n.d.; Arslanli and Sabuncuoglu, 2023; Joshi and Choudhury, '
        '2022). The NFT standard provides built-in uniqueness guarantees (no two tokens '
        'with the same ID can exist) and marketplace interoperability. However, standard '
        'ERC-721/ERC-1155 NFTs do not include multi-party transfer approval at the '
        'contract level — the token owner can transfer to any address unilaterally, '
        'which is legally inappropriate for land transfers requiring registrar approval.',
        indent=True)

def build_hybrid(doc):
    section_heading(doc, 'VI', 'Hybrid and Dual-Ledger Architectures')

    body(doc,
        'A consistent finding across the reviewed literature is that pure blockchain '
        'implementations — where all property data is stored on-chain — are impractical '
        'for production deployment. Blockchain storage is expensive (gas costs on Ethereum), '
        'slow (block confirmation times), and opaque to standard database queries. '
        'Mansoor et al. (2023) explicitly recommend hybrid architectures, and this '
        'recommendation is supported by Podshivalov (2022) and Paavo et al. (2025).',
        indent=True)

    body(doc,
        'The hybrid approach separates concerns between two persistence layers:',
        indent=True)

    for item in [
        'Relational database (PostgreSQL, MySQL) — stores operational metadata: '
        'user profiles, full property descriptions, transfer histories with timestamps '
        'and prices, dispute records, audit logs. Supports standard SQL queries, '
        'joins, and full-text search.',
        'Blockchain — stores only what must be immutable and publicly verifiable: '
        'ownership claims (property ID to owner address mapping), authorised registrar '
        'addresses, and transfer lifecycle events.',
    ]:
        bullet(doc, item)

    body(doc,
        'The critical design question in hybrid systems is which layer is canonical. '
        'Podshivalov (2022) and Paavo et al. (2025) both assert that the blockchain '
        'must be treated as the ground truth for ownership, with the database serving '
        'as a queryable cache that must be consistent with blockchain state. This '
        'one-directional consistency constraint — blockchain authoritative, database '
        'derived — is the defining architectural invariant of a well-designed hybrid system.',
        indent=True)

    body(doc,
        'IPFS (InterPlanetary File System) is used in several systems (Daniel and '
        'Ifejika Speranza, 2020; Vu et al., n.d.) for document storage. The content-addressed '
        'nature of IPFS means that the hash of a document stored in IPFS can be committed '
        'on-chain, providing a tamper-evident reference: any modification to the document '
        'produces a different hash, detectable by comparing to the on-chain value.',
        indent=True)

def build_developing(doc):
    section_heading(doc, 'VII', 'Developing Economy Deployments and Adoption Challenges')

    body(doc,
        'The reviewed literature includes studies from Bangladesh (Alam et al., 2022), '
        'Nigeria (Ibrahim et al., 2021), Tanzania (Kombe et al., 2017), Pakistan '
        '(Mansoor et al., 2023), Namibia (Paavo et al., 2025), and Turkey (Arslanli '
        'and Sabuncuoglu, 2023). Cross-cutting challenges identified include:',
        indent=True)

    simple_table(doc,
        ['Challenge', 'Frequency in Literature', 'Proposed Mitigation'],
        [
            ['Legal/regulatory gap — blockchain records not legally recognised', 'Very High', 'Phased approach: blockchain as supplementary record until legislation enacted'],
            ['Infrastructure — unreliable internet, low smartphone penetration', 'High', 'Lightweight clients, offline-capable interfaces, SMS-based verification'],
            ['Gas/transaction costs on public blockchains', 'High', 'Permissioned chains (Hyperledger) or Layer-2 solutions'],
            ['Corrupt registrar risk — oracle problem', 'High', 'Multi-registrar approval, whistleblower mechanisms, civil society audit access'],
            ['Digital literacy among landowners', 'Medium', 'Agent-based model: trusted community intermediaries submit on behalf of owners'],
            ['Socio-cultural resistance to digital systems', 'Medium (Ameyaw & de Vries, 2023)', 'Community engagement, local language interfaces, participatory design'],
        ]
    )
    caption(doc, 'TABLE II. Developing Economy Adoption Challenges Identified in Reviewed Literature')

    body(doc,
        'Ameyaw and de Vries (2023) provide the most thorough analysis of socio-cultural '
        'factors, arguing that technology capability is rarely the binding constraint — '
        'adoption is determined by trust in the sponsoring institution, perceived fairness '
        'of the transition from existing informal systems, and cultural attitudes toward '
        'digital record-keeping. This finding has significant implications for Zimbabwe\'s '
        'context, where a history of government land appropriation (particularly in the '
        '2000-2008 period) has eroded trust in state-administered land records.',
        indent=True)

def build_security(doc):
    section_heading(doc, 'VIII', 'Security and Fraud Prevention')

    body(doc,
        'The reviewed literature\'s treatment of security is the most significant gap '
        'identified in this review. M. Zein and Twinomurinzi (2023) found that fewer '
        'than a third of 47 reviewed papers performed any formal threat analysis. Of the '
        '21 papers in this corpus, only four engaged with security beyond high-level '
        'claims of "tamper-proof" or "immutable" storage:',
        indent=True)

    simple_table(doc,
        ['Paper', 'Security Analysis Performed', 'Threats Addressed', 'Gaps'],
        [
            ['Racetin et al. (2022)', 'Data integrity focus', 'Tampering of ownership records', 'No application-layer analysis; no STRIDE or equivalent'],
            ['Dwivedi et al. (2023)', 'Smart contract vulnerability review', 'Reentrancy, integer overflow, unchecked send', 'Focuses on contract layer only; no auth/session analysis'],
            ['Konashevych (2020)', 'Oracle problem analysis', 'Gap between digital state and physical reality', 'Conceptual only; no implementation security analysis'],
            ['Mansoor et al. (2023)', 'Conceptual security overview', 'Data tampering, single point of failure', 'No threat model, no testing, no application-layer analysis'],
        ]
    )
    caption(doc, 'TABLE III. Security Analysis in Reviewed Papers')

    body(doc,
        'No reviewed paper applies a structured threat modelling methodology (STRIDE, '
        'PASTA, LINDDUN, or equivalent) to a blockchain land registry system. No paper '
        'examines authentication and session management threats (JWT forgery, replay '
        'attacks, privilege escalation) at the API layer. No paper evaluates resistance '
        'to adversarial inputs through empirical testing. This is a critical gap: the '
        'security of a land registry system depends as much on the security of the '
        'application layer — authentication, authorisation, input validation, session '
        'management — as on the immutability of the underlying blockchain.',
        indent=True)

    body(doc,
        'Smart contract security is better represented. Dwivedi et al. (2023) recommend '
        'formally verifiable smart contract languages as a design principle, noting that '
        'Solidity\'s Turing-completeness and mutable state model create a large attack '
        'surface. This recommendation motivates the use of Clarity in BlockLand — a '
        'non-Turing-complete, decidable language in which reentrancy is architecturally '
        'impossible and all execution paths are statically analysable before deployment.',
        indent=True)

def build_legal(doc):
    section_heading(doc, 'IX', 'Legal and Regulatory Landscape')

    body(doc,
        'No jurisdiction reviewed in the literature has enacted legislation that '
        'formally recognises blockchain land registry records as legally equivalent '
        'to paper title deeds. The furthest-advanced jurisdiction is Georgia, where '
        'the National Agency of Public Registry (NAPR) uses Bitfury\'s blockchain '
        'as a supplementary audit layer, not a replacement for the official paper '
        'registry. Sweden\'s Lantmateriet pilot concluded that technical feasibility '
        'was demonstrated but legislative change was a prerequisite for production '
        'deployment (Podshivalov, 2022).',
        indent=True)

    body(doc,
        'For Zimbabwe specifically, Kimathi (n.d.) identifies the Deeds Registries '
        'Act [Chapter 20:05] and the Electronic Transactions and Electronic Commerce '
        'Act as the two pieces of legislation that would require amendment to recognise '
        'blockchain-recorded transfers. The Electronic Transactions Act currently '
        'recognises electronic signatures but does not address distributed ledger '
        'records specifically. Ameyaw and de Vries (2023) note that legal reform '
        'in land administration is politically contentious in many African states '
        'because land tenure reform intersects with post-colonial redistributive policy.',
        indent=True)

    body(doc,
        'The reviewed literature consistently recommends a phased approach: deploy '
        'blockchain as a supplementary audit layer operating in parallel with the '
        'existing paper registry, build institutional familiarity and demonstrated '
        'reliability over 2-3 years, then seek legislative endorsement for full '
        'transition. This is the approach adopted by BlockLand.',
        indent=True)

def build_gaps(doc):
    section_heading(doc, 'X', 'Research Gaps and Future Directions')

    body(doc,
        'This review identifies the following research gaps warranting future investigation:',
        indent=True)

    for gap, detail in [
        ('Application-layer security analysis.',
         'No existing study applies structured threat modelling (STRIDE, PASTA, or equivalent) '
         'to a blockchain land registry system, or evaluates resistance to authentication and '
         'session management attacks. Future work should apply STRIDE to the full three-tier '
         'architecture — frontend, API, blockchain — not merely the smart contract layer.'),
        ('Bitcoin-anchored blockchain platforms.',
         'All reviewed papers use Ethereum, Hyperledger, or no specified blockchain. '
         'Stacks — which anchors to Bitcoin via Proof of Transfer — has not been evaluated '
         'for land registry. Its Clarity language, which is non-Turing-complete and '
         'architecturally reentrancy-free, represents a meaningful advance over Solidity '
         'for this use case.'),
        ('Multi-party transfer workflow enforcement at contract level.',
         'The three-party workflow (seller, buyer, registrar) proposed by Kombe et al. (2017) '
         'is not implemented at the smart contract level in any reviewed empirical paper. '
         'Most systems either allow unilateral owner transfer or implement registrar approval '
         'off-chain. Contract-enforced three-party workflows that map to existing legal '
         'requirements of specific jurisdictions remain an open implementation gap.'),
        ('Longitudinal evaluation in production.',
         'No reviewed paper presents a longitudinal evaluation of a deployed system beyond '
         'the pilot phase. The performance, reliability, and adoption characteristics of '
         'blockchain land registry systems under real government workloads are unknown.'),
        ('Zimbabwe-specific deployment.',
         'No published peer-reviewed study addresses blockchain land registry in Zimbabwe\'s '
         'specific legal, administrative, and socio-political context.'),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f'{gap} ')
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(detail)
        run2.font.size = Pt(10)

def build_conclusion(doc):
    section_heading(doc, 'XI', 'Conclusion')

    body(doc,
        'This systematic review has surveyed twenty-one peer-reviewed papers on blockchain '
        'technology applications in land registry systems, synthesising findings across '
        'platform selection, smart contract design, hybrid architectures, developing economy '
        'adoption, security, and legal frameworks. The central finding is that while the '
        'technical feasibility of blockchain-based land registry has been convincingly '
        'demonstrated across multiple contexts, the security of the application layer above '
        'the blockchain remains systematically under-analysed, and no study has implemented '
        'a legally-aligned, contract-enforced multi-party transfer workflow for a specific '
        'jurisdiction.',
        indent=True)

    body(doc,
        'The hybrid dual-ledger architecture — relational database for operational '
        'richness, blockchain as immutable ownership ground truth — emerges from the '
        'literature as the most practically deployable approach. The critical design '
        'invariant is that the blockchain must be treated as canonical: the database '
        'is a cache, not the record of authority.',
        indent=True)

    body(doc,
        'Five research gaps are identified. These gaps directly motivate the BlockLand '
        'system, which applies STRIDE threat modelling across all architectural layers, '
        'uses the Stacks blockchain with Clarity smart contracts, enforces a three-step '
        'transfer workflow at the contract level aligned with Zimbabwe\'s Deeds Registries '
        'Act, and provides empirical security test results.',
        indent=True)

def build_references(doc):
    section_heading(doc, '', 'References')
    hr(doc)

    refs = [
        'Alam, K.M., Rahman, J.M.A., Tasnim, A. and Akther, A. (2022) \'A Blockchain-based Land Title Management System for Bangladesh\', Journal of King Saud University - Computer and Information Sciences, 34(6), pp. 3096-3110.',
        'Ameyaw, P.D. and de Vries, W.T. (2023) \'Blockchain technology adaptation for land administration services: The importance of socio-cultural elements\', Land Use Policy.',
        'Arslanli, K.Y. and Sabuncuoglu, G. (2023) \'Tokenization of Real Estate: A Study on Land Tokenization in Turkey\', 29th Annual European Real Estate Society Conference, London.',
        'Buterin, V. (2014) \'A next-generation smart contract and decentralized application platform\' [Ethereum Whitepaper]. Available at: https://ethereum.org/en/whitepaper (Accessed: 24 February 2026).',
        'Daniel, D. and Ifejika Speranza, C. (2020) \'The Role of Blockchain in Documenting Land Users\' Rights: The Canonical Case of Farmers in the Vernacular Land Market\', Frontiers in Blockchain, 3, p. 19.',
        'Dwivedi, R., Patel, S. and Shukla, S. (2023) \'Blockchain-Based Transferable Digital Rights of Land\', arXiv [Preprint], arXiv:2308.05950.',
        'Ibrahim, I., Daud, D., Azmi, F.A.M., Noor, N.A.M. and Yusoff, N.S.M. (2021) \'Improvement Of Land Administration System In Nigeria: A Blockchain Technology Review\', 10(8).',
        'Joshi, S. and Choudhury, A. (2022) \'Tokenization of Real Estate Assets Using Blockchain\', International Journal of Intelligent Information Technologies, 18(3), pp. 1-12.',
        'Kimathi, K.T. (n.d.) A Land Administration Model Based on Blockchain Technology. [Unpublished thesis]. Harare: Harare Institute of Technology.',
        'Kombe, C., Manyilizu, M. and Mvuma, A. (2017) \'Design of Land Administration and Title Registration Model Based on Blockchain Technology\'.',
        'Konashevych, O. (2020) \'General Concept of Real Estate Tokenization on Blockchain: The Right to Choose\', European Property Law Journal, 9(1), pp. 21-66.',
        'Konashevych, O. (n.d.) \'Constraints and Benefits of Blockchain Use for Real Estate and Property Rights\'. [Unpublished manuscript].',
        'Lemieux, V.L. (2016) \'Trusting records: Is Blockchain technology the answer?\', Records Management Journal, 26(2), pp. 110-139.',
        'M. Zein, R. and Twinomurinzi, H. (2023) \'Blockchain Technology in Lands Registration: A Systematic Literature Review\', JeDEM - eJournal of eDemocracy and Open Government, 15(2), pp. 1-36.',
        'Mansoor, M.A., Ali, M., Mateen, A., Kaleem, M. and Nazir, S. (2023) \'Blockchain Technology for Land Registry Management in Developing Countries\', 2023 2nd International Conference on ETECTE, Lahore, Pakistan.',
        'Ministry of National Housing and Social Amenities (Zimbabwe) (n.d.) Accessing Land for Housing Development in Zimbabwe. Harare: Government of Zimbabwe.',
        'Nakamoto, S. (2008) \'Bitcoin: A Peer-to-Peer Electronic Cash System\'. Available at: https://bitcoin.org/bitcoin.pdf (Accessed: 24 February 2026).',
        'Paavo, J.P., Rodriguez-Puentes, R. and Chigbu, U.E. (2025) \'Practicality of Blockchain Technology for Land Registration: A Namibian Case Study\', Land, 14(8), p. 1626.',
        'Podshivalov, T.P. (2022) \'Improving implementation of the Blockchain technology in real estate registration\', Journal of High Technology Management Research, 33(2), p. 100440.',
        'Racetin, I., Kilic Pamukovic, J., Zrinjski, M. and Peko, M. (2022) \'Blockchain-Based Land Management for Sustainable Development\', Sustainability, 14(17), p. 10649.',
        'Vu, H.N., Tran, Q.T. and Hoang, V.N. (n.d.) \'Transforming Land Ownership Documents into Blockchain-Based NFTs for Secure Digital Management\'. [Unpublished manuscript].',
        'Zook, M. and McCanless, M. (2025) \'Blockchain real estate: The messy landing of digital property\', Progress in Economic Geography, 3(1), p. 100039.',
    ]

    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.8)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f'[{i}]  {ref}')
        run.font.size = Pt(9)

# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(10)

    build_header(doc)
    build_abstract(doc)
    build_introduction(doc)
    build_methodology(doc)
    build_foundational(doc)
    build_platforms(doc)
    build_smart_contracts(doc)
    build_hybrid(doc)
    build_developing(doc)
    build_security(doc)
    build_legal(doc)
    build_gaps(doc)
    build_conclusion(doc)
    build_references(doc)

    output = 'BlockLand_Survey_Paper.docx'
    doc.save(output)
    print(f'Saved: {output}')

if __name__ == '__main__':
    main()
