"""
Generate BlockLand Project Proposal Word document.
Run: python generate_proposal.py
Output: BlockLand_Project_Proposal.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK   = RGBColor(0x00, 0x00, 0x00)
NAVY    = RGBColor(0x0F, 0x17, 0x2A)
TEAL    = RGBColor(0x0D, 0x94, 0x88)
GRAY    = RGBColor(0x44, 0x44, 0x44)
RED_HIT = RGBColor(0xC0, 0x00, 0x00)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def hr(doc, color='0D9488'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = NAVY
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '0D9488')
    pBdr.append(bot)
    pPr.append(pBdr)

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = TEAL

def body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK

def bullet(doc, text, level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)
    run = p.add_run(f'•  {text}')
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK

def numbered_item(doc, number, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run_n = p.add_run(f'{number}. ')
    run_n.bold = True
    run_n.font.size = Pt(11)
    run_n.font.color.rgb = TEAL
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK

def info_row(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run_l = p.add_run(f'{label}: ')
    run_l.bold = True
    run_l.font.size = Pt(11)
    run_l.font.color.rgb = NAVY
    run_v = p.add_run(value)
    run_v.font.size = Pt(11)
    run_v.font.color.rgb = BLACK

def simple_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], '0F172A')
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        bg = 'EFF6FF' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            set_cell_bg(cells[ci], bg)
            p = cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
    doc.add_paragraph()

# ── Cover page ────────────────────────────────────────────────────────────────

def build_cover(doc):
    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('HARARE INSTITUTE OF TECHNOLOGY')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_HIT
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Department of Information Security and Assurance')
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Bachelor of Technology in Information Security and Assurance')
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = GRAY
    p.paragraph_format.space_after = Pt(24)

    hr(doc, 'C00000')

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PROJECT PROPOSAL')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BlockLand: A Blockchain-Based Land Registry and\nManagement System for Zimbabwe')
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RED_HIT
    p.paragraph_format.space_after = Pt(32)

    for _ in range(2):
        doc.add_paragraph()

    for label, value in [
        ('Student Name', 'Tinotenda James'),
        ('Student Number', 'H220349M'),
        ('Supervisor', 'Ms Chavhunduka'),
        ('Department', 'Information Security and Assurance'),
        ('Date of Submission', '3 September 2025'),
    ]:
        info_row(doc, label, value)

    for _ in range(3):
        doc.add_paragraph()
    hr(doc, 'C00000')

# ── Sections ──────────────────────────────────────────────────────────────────

def build_background(doc):
    doc.add_page_break()
    heading1(doc, '1. Introduction and Background')
    body(doc,
        'Land is the most economically significant asset class in Zimbabwe, yet the '
        'administrative systems that record, verify, and transfer ownership remain '
        'predominantly paper-based and centralised within the Deeds Registry, operating '
        'under the Deeds Registries Act [Chapter 20:05]. Physical title deed documents, '
        'manual ledgers, and in-person approval processes create systemic vulnerabilities: '
        'title deeds can be forged, ledger entries altered, and ownership disputed without '
        'a tamper-evident historical record.',
        indent=True)

    body(doc,
        'Blockchain technology offers a class of properties particularly suited to '
        'addressing these vulnerabilities: immutability of the historical ownership record, '
        'transparent auditability without reliance on a single trusted authority, and '
        'programmatic enforcement of multi-party workflows through smart contracts. '
        'Several international pilots — in Georgia, Ghana, Sweden, and Namibia — have '
        'demonstrated the technical feasibility of blockchain-based land registry, though '
        'no production system tailored to Zimbabwe\'s legal and administrative context '
        'currently exists.',
        indent=True)

    body(doc,
        'This proposal describes BlockLand: a dual-ledger land registry system that '
        'combines the operational richness of a relational database with the immutability '
        'and transparency of the Stacks blockchain — a Layer-1 blockchain that anchors '
        'its consensus to the Bitcoin network via Proof of Transfer (PoX).',
        indent=True)

def build_problem(doc):
    heading1(doc, '2. Problem Statement')
    body(doc,
        'The current land administration process in Zimbabwe is characterised by five '
        'interrelated problems:',
        indent=True)
    for item in [
        'Title deed forgery — physical documents can be duplicated or falsified without a cryptographic reference.',
        'Opacity of ownership history — there is no publicly accessible, authoritative record of who owned a property and when.',
        'Single-point-of-failure — the Deeds Registry is a centralised record-keeper whose integrity cannot be independently verified.',
        'Slow, paper-dependent transfer workflow — property transfers require physical document exchange, in-person visits, and manual stamp approvals that take weeks.',
        'No public verification mechanism — citizens, banks, and lawyers cannot instantly verify ownership without formal inquiry to the Registry.',
    ]:
        bullet(doc, item)

    body(doc,
        'These problems create fertile ground for land fraud, disputed ownership claims, '
        'and a general erosion of public confidence in the land administration system. '
        'BlockLand is proposed as a targeted technical intervention to address all five '
        'problems through a combination of blockchain anchoring and a digital workflow '
        'system.',
        indent=True)

def build_objectives(doc):
    heading1(doc, '3. Aim and Objectives')
    heading2(doc, '3.1  Project Aim')
    body(doc,
        'To design, develop, and evaluate a blockchain-based land registry system for '
        'Zimbabwe that eliminates paper-based title deed fraud, provides transparent '
        'ownership verification, and enforces a legally-aligned multi-party transfer '
        'workflow on an immutable, publicly auditable ledger.',
        indent=True)

    heading2(doc, '3.2  Specific Objectives')
    numbered_item(doc, 1,
        'To provide a digital submission, review, and approval workflow for property '
        'title registrations — eliminating paper-based processes and reducing registration '
        'time from weeks to minutes.')
    numbered_item(doc, 2,
        'To assign each property a unique on-chain token identifier, making it '
        'cryptographically impossible to register the same plot twice under different '
        'ownership records.')
    numbered_item(doc, 3,
        'To expose a public verification portal where any citizen can instantly confirm '
        'property ownership by plot number, deed number, or owner identification — '
        'without requiring a registered account or government inquiry.')

def build_solution(doc):
    heading1(doc, '4. Proposed Solution')
    body(doc,
        'BlockLand is a three-tier web application comprising a Next.js 14 frontend, '
        'a NestJS 10 REST API backend, and a dual persistence layer:',
        indent=True)

    bullet(doc, 'PostgreSQL — stores all operational metadata: user profiles, property details, '
           'marketplace listings, transfer histories with sale prices, payment records, disputes, '
           'and a full audit log.')
    bullet(doc, 'Stacks Blockchain — stores only what must be immutable: property ownership '
           'assertions (plot-to-owner mappings), authorised registrar wallet addresses, and '
           'the three-step transfer state machine.')

    body(doc,
        'A Clarity smart contract enforces the three-step ownership transfer workflow: '
        '(1) the current owner initiates the transfer, naming the buyer; (2) the buyer '
        'signs their approval on-chain; (3) an authorised registrar finalises the transfer, '
        'atomically reassigning ownership. No single party can complete a transfer '
        'unilaterally. IPFS is used for decentralised title deed document storage, with '
        'the content-addressed hash committed on-chain to prevent silent document substitution.',
        indent=True)

    body(doc,
        'The system implements JWT RS256 dual-token authentication and a four-role '
        'RBAC model (Admin, Registrar, Owner, Buyer). A public /verify portal requires '
        'no login and answers ownership queries in real time by querying the on-chain '
        'state.',
        indent=True)

def build_scope(doc):
    heading1(doc, '5. Scope and Limitations')
    heading2(doc, '5.1  In Scope')
    for item in [
        'Property registration workflow (digital title deed submission, registrar review and approval, on-chain registration).',
        'Three-step ownership transfer workflow enforced at both API and smart contract layers.',
        'Public verification portal (no account required) — search by plot number, deed number, or owner National ID.',
        'Marketplace module for listing properties and expressing buyer interest.',
        'Dispute flagging and registrar resolution workflow.',
        'PDF title deed certificate generation on demand.',
        'Role-based access control: Admin, Registrar, Owner, Buyer.',
        'JWT RS256 authentication with access and refresh token model.',
    ]:
        bullet(doc, item)

    heading2(doc, '5.2  Out of Scope')
    for item in [
        'Integration with the official Zimbabwe Deeds Registry backend systems (legal mandate not obtained for prototype phase).',
        'Mobile application (web browser only for this iteration).',
        'Mainnet deployment — the system operates on the Stacks testnet throughout development.',
        'Automated valuation or taxation modules.',
        'Physical boundary survey data or GIS map integration.',
    ]:
        bullet(doc, item)

def build_methodology(doc):
    heading1(doc, '6. Development Methodology')
    body(doc,
        'The Waterfall methodology has been selected for this project. BlockLand has '
        'well-defined, stable requirements rooted in the existing legal framework of the '
        'Zimbabwe Deeds Registries Act, and the multi-party nature of the blockchain '
        'workflow demands that each phase be fully completed and verified before the next '
        'begins. The sequential, document-driven nature of Waterfall aligns with the '
        'academic reporting requirements of the HIT final-year project.',
        indent=True)

    simple_table(doc,
        ['Phase', 'Activity', 'Duration'],
        [
            ['1. Requirements Analysis', 'Stakeholder interviews, legal review, system requirements specification', 'Sept - Oct 2025'],
            ['2. System Design', 'Architecture design, database schema, smart contract specification, UI wireframes', 'Oct - Nov 2025'],
            ['3. Implementation', 'Smart contract development (Clarity), backend (NestJS), frontend (Next.js)', 'Nov 2025 - Feb 2026'],
            ['4. Testing', 'Unit testing, integration testing, adversarial security testing, user acceptance testing', 'Feb - Mar 2026'],
            ['5. Documentation', 'Dissertation write-up, user manual, technical documentation', 'Mar - Apr 2026'],
            ['6. Submission', 'Final submission and viva voce', 'April 2026'],
        ]
    )

def build_gantt(doc):
    heading1(doc, '7. Project Timeline (Gantt Chart)')

    months = ['Sep\n25', 'Oct\n25', 'Nov\n25', 'Dec\n25', 'Jan\n26', 'Feb\n26', 'Mar\n26', 'Apr\n26']
    tasks = [
        ('Requirements Analysis',  [1,1,0,0,0,0,0,0]),
        ('System Design',          [0,1,1,0,0,0,0,0]),
        ('Smart Contract Dev.',    [0,0,1,1,0,0,0,0]),
        ('Backend Development',    [0,0,1,1,1,0,0,0]),
        ('Frontend Development',   [0,0,0,1,1,1,0,0]),
        ('Testing',                [0,0,0,0,0,1,1,0]),
        ('Documentation',          [0,0,0,0,0,0,1,1]),
        ('Submission',             [0,0,0,0,0,0,0,1]),
    ]

    table = doc.add_table(rows=1 + len(tasks), cols=1 + len(months))
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    set_cell_bg(hdr[0], '0F172A')
    run = hdr[0].paragraphs[0].add_run('Task')
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, m in enumerate(months):
        set_cell_bg(hdr[i+1], '0F172A')
        p = hdr[i+1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(m)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for ri, (task, marks) in enumerate(tasks):
        cells = table.rows[ri+1].cells
        set_cell_bg(cells[0], 'F8FAFC')
        run = cells[0].paragraphs[0].add_run(task)
        run.font.size = Pt(8)
        for ci, m in enumerate(marks):
            if m:
                set_cell_bg(cells[ci+1], '0D9488')
                p = cells[ci+1].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run('')
                run.font.size = Pt(8)
            else:
                set_cell_bg(cells[ci+1], 'FFFFFF')

    doc.add_paragraph()

def build_resources(doc):
    heading1(doc, '8. Resources Required')
    simple_table(doc,
        ['Resource', 'Specification', 'Purpose'],
        [
            ['Development Machine', 'Windows 11, 16 GB RAM, Node.js 20', 'All development and testing'],
            ['PostgreSQL 15', 'Local and cloud-hosted instance', 'Primary relational database'],
            ['Stacks Testnet / Clarinet', 'Clarinet v2, Stacks devnet', 'Smart contract development and testing'],
            ['IPFS Node (Helia / Web3.Storage)', 'Free tier sufficient', 'Decentralised document storage'],
            ['Vercel / Railway', 'Free tier for prototype deployment', 'Frontend and backend hosting'],
            ['GitHub', 'Private repository', 'Version control and CI/CD'],
            ['Hiro Wallet (browser extension)', 'Free, testnet mode', 'Blockchain transaction signing'],
        ]
    )

def build_deliverables(doc):
    heading1(doc, '9. Expected Deliverables')
    for item in [
        'Fully functional BlockLand web application (frontend + backend + smart contract).',
        'Deployed Clarity smart contract on Stacks testnet with verified source code.',
        'PostgreSQL database schema and seed data.',
        'Security test report covering STRIDE threat analysis and adversarial test results.',
        'User manual for all four system roles (Admin, Registrar, Owner, Buyer).',
        'Final-year dissertation (seven chapters, including full system documentation).',
        'Research paper and technical paper as required academic outputs.',
    ]:
        bullet(doc, item)

def build_references(doc):
    heading1(doc, 'References')
    for ref in [
        'M. Zein, R. and Twinomurinzi, H. (2023) \'Blockchain Technology in Lands Registration: A Systematic Literature Review\', JeDEM - eJournal of eDemocracy and Open Government, 15(2), pp. 1-36.',
        'Mansoor, M.A., Ali, M., Mateen, A., Kaleem, M. and Nazir, S. (2023) \'Blockchain Technology for Land Registry Management in Developing Countries\', 2023 2nd International Conference on Emerging Trends in Electrical, Control, and Telecommunication Engineering (ETECTE), Lahore, Pakistan.',
        'Ministry of National Housing and Social Amenities (Zimbabwe) (n.d.) Accessing Land for Housing Development in Zimbabwe. Harare: Government of Zimbabwe.',
        'Paavo, J.P., Rodriguez-Puentes, R. and Chigbu, U.E. (2025) \'Practicality of Blockchain Technology for Land Registration: A Namibian Case Study\', Land, 14(8), p. 1626.',
        'Kombe, C., Manyilizu, M. and Mvuma, A. (2017) \'Design of Land Administration and Title Registration Model Based on Blockchain Technology\'.',
        'Kimathi, K.T. (n.d.) A Land Administration Model Based on Blockchain Technology. [Unpublished thesis]. Harare: Harare Institute of Technology.',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(ref)
        run.font.size = Pt(10)

# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(11)

    build_cover(doc)
    build_background(doc)
    build_problem(doc)
    build_objectives(doc)
    build_solution(doc)
    build_scope(doc)
    build_methodology(doc)
    build_gantt(doc)
    build_resources(doc)
    build_deliverables(doc)
    build_references(doc)

    output = 'BlockLand_Project_Proposal.docx'
    doc.save(output)
    print(f'Saved: {output}')

if __name__ == '__main__':
    main()
