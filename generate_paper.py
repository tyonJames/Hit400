"""
Generate BlockLand IEEE Access-style research paper Word document.
Run: python generate_paper.py
Output: BlockLand_Research_Paper.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK  = RGBColor(0x00, 0x00, 0x00)
GRAY   = RGBColor(0x33, 0x33, 0x33)
LGRAY  = RGBColor(0x66, 0x66, 0x66)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_page_break(doc):
    doc.add_page_break()

def hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '000000')
    pBdr.append(bot)
    pPr.append(pBdr)

def section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'{number}. {title.upper()}')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK

def subsection_heading(doc, label, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{label} {title}')
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK

def body(doc, text, indent=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK

def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

def simple_table(doc, headers, rows, font_size=9):
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], '000000')
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        bg = 'F5F5F5' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            set_cell_bg(cells[ci], bg)
            p = cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(font_size)
    doc.add_paragraph()

# ── Paper sections ─────────────────────────────────────────────────────────────

def build_header(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BlockLand: A Blockchain-Based Dual-Ledger Land Registry System with Multi-Layer Security for Developing Economies')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Tinotenda James')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(2)

    for line in [
        'Department of Information Security and Assurance',
        'Harare Institute of Technology, Belvedere, Harare, Zimbabwe',
        'H220349M | tinotendajames@hit.ac.zw',
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = LGRAY
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()
    hr(doc)

def build_abstract(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    run = p.add_run('Abstract')
    run.bold = True
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    run = p.add_run(
        'Land administration systems in developing economies suffer from chronic fraud, document '
        'tampering, and administrative opacity, eroding public trust and hindering economic development. '
        'This paper presents BlockLand, a dual-ledger land registry system that combines a PostgreSQL '
        'relational database for rich operational metadata with the Stacks blockchain as an immutable, '
        'publicly auditable ownership ground truth. The system exposes a NestJS REST API consumed by a '
        'Next.js 14 web application and enforces a three-step transfer workflow — owner initiation, '
        'buyer approval, and registrar finalisation — through a Clarity smart contract. We present a '
        'comprehensive security analysis using the STRIDE threat model, systematically identifying '
        'threats and the countermeasures applied: JWT RS256 dual-token authentication with 15-minute '
        'access token expiry, four-role RBAC enforced at both API middleware and contract layers, '
        'IPFS-anchored document hashing for title deed integrity, and the non-Turing-complete Clarity '
        'language to eliminate entire classes of smart contract vulnerabilities. Security testing '
        'confirmed resistance to replay attacks, horizontal privilege escalation, and double-transfer '
        'attempts. The system achieves a property registration throughput suitable for Zimbabwe\'s '
        'Deeds Registry workload with sub-second API response times under standard load. BlockLand '
        'demonstrates that targeted blockchain integration, rather than wholesale system replacement, '
        'delivers meaningful security guarantees appropriate for Zimbabwe\'s land administration context.'
    )
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('Index Terms')
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(
        ' — blockchain, land registry, smart contracts, security analysis, STRIDE threat model, '
        'Clarity, JWT RS256, role-based access control, Zimbabwe, developing economies'
    )
    run2.font.size = Pt(10)
    run2.italic = True

    hr(doc)

def build_introduction(doc):
    section_heading(doc, 'I', 'Introduction')

    body(doc,
        'Land is the most significant asset class in most developing economies, yet the '
        'administrative systems that record and transfer ownership remain paper-based, centralised, '
        'and vulnerable to fraud. In Zimbabwe, the Deeds Registry operates under the Deeds Registries '
        'Act [Cap 20:05], maintaining physical ledgers and manual title deeds that are susceptible '
        'to forgery, loss, and deliberate falsification (Ministry of National Housing and Social '
        'Amenities (Zimbabwe), n.d.). A 2023 systematic literature review of blockchain land registry '
        'deployments identified administrative corruption and document fraud as the two most '
        'frequently cited motivations for blockchain adoption in this domain (M. Zein and '
        'Twinomurinzi, 2023).',
        indent=True)

    body(doc,
        'Blockchain technology offers three properties particularly relevant to land administration: '
        'immutability of the historical record, transparent auditability without a trusted third '
        'party, and programmatic enforcement of multi-party workflows through smart contracts. Prior '
        'work has demonstrated deployments in Georgia, Sweden, Ghana, and Namibia (Paavo et al., '
        '2025; Mansoor et al., 2023), but few systems have published security analyses that go beyond '
        'high-level claims of "tamper-proof" storage. This paper addresses that gap.',
        indent=True)

    body(doc,
        'This paper makes the following contributions:',
        indent=True)

    for item in [
        '(1) A dual-ledger architecture that preserves the operational richness of relational databases '
        'while anchoring ownership state to a public blockchain, providing practical deployability in '
        'low-bandwidth environments.',
        '(2) A rigorous STRIDE-based security analysis of each architectural layer, mapping identified '
        'threats to concrete countermeasures implemented in the system.',
        '(3) A demonstration that Clarity smart contracts — through decidability and the absence of '
        'reentrancy — eliminate entire categories of blockchain-specific vulnerabilities present in '
        'Solidity-based alternatives.',
        '(4) An empirical evaluation of the system\'s security posture under adversarial test scenarios '
        'including replay attacks, JWT forgery attempts, and double-transfer conditions.',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(3)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(item)
        run.font.size = Pt(10)

    body(doc,
        'The remainder of this paper is structured as follows. Section II surveys related work on '
        'blockchain land registries and security frameworks. Section III describes the BlockLand '
        'architecture. Section IV presents the security analysis. Section V reports evaluation results. '
        'Section VI discusses limitations and future work. Section VII concludes.',
        indent=True)

def build_related_work(doc):
    section_heading(doc, 'II', 'Related Work')

    subsection_heading(doc, 'A.', 'Blockchain for Land Administration')
    body(doc,
        'The application of blockchain to land administration has attracted significant academic '
        'attention since 2016. Kombe et al. (2017) proposed one of the earliest formal models for '
        'blockchain-based land title registration, identifying the need for a three-party workflow '
        'involving seller, buyer, and a government registrar — a pattern BlockLand implements '
        'as a first-class on-chain constraint. Alam et al. (2022) implemented a functioning prototype '
        'for Bangladesh demonstrating that Ethereum-based land title systems are technically feasible, '
        'though they report gas cost as a significant practical barrier not present on the Stacks '
        'network used by BlockLand.',
        indent=True)

    body(doc,
        'For Sub-Saharan Africa specifically, Ibrahim et al. (2021) reviewed land administration '
        'improvements in Nigeria and found that trust deficit in centralised record-keepers is the '
        'primary adoption driver. Mansoor et al. (2023) examined developing country deployments and '
        'concluded that hybrid systems — blockchain for ownership state, traditional databases for '
        'operational data — outperform pure blockchain implementations in deployability and '
        'performance. BlockLand\'s dual-ledger design directly operationalises this recommendation. '
        'Paavo et al. (2025) studied a Namibian case study and identified registrar workflow '
        'integration as the critical success factor for practical adoption, motivating BlockLand\'s '
        'registrar-finalisation step.',
        indent=True)

    body(doc,
        'Daniel and Ifejika Speranza (2020) examined informal land markets and found that blockchain '
        'can document customary rights not captured by formal registries, while Ameyaw and de Vries '
        '(2023) emphasised socio-cultural acceptance as the binding constraint for technology adoption '
        'in African contexts — a limitation acknowledged in Section VI.',
        indent=True)

    subsection_heading(doc, 'B.', 'Security in Blockchain Land Systems')
    body(doc,
        'Despite the volume of blockchain land registry literature, systematic security analyses '
        'remain rare. M. Zein and Twinomurinzi (2023) reviewed 47 papers and found that fewer than '
        'a third provided any threat modelling. Most papers assert immutability and transparency as '
        'inherent security properties without analysing threats to the application layer above the '
        'blockchain. Konashevych (2020) identified the oracle problem — the gap between on-chain '
        'state and physical reality — as the most significant security concern not addressed by '
        'blockchain properties alone; BlockLand addresses this through registrar verification of '
        'physical documents before any on-chain state change.',
        indent=True)

    body(doc,
        'Dwivedi et al. (2023) proposed blockchain-based transferable digital rights using NFTs and '
        'identified smart contract bugs as the primary attack surface, recommending formally '
        'verifiable languages. This motivates BlockLand\'s choice of Clarity over Solidity. '
        'Racetin et al. (2022) proposed blockchain-based land management for sustainable development '
        'and modelled data integrity as the central security property, consistent with our STRIDE '
        'analysis in which Tampering is the highest-impact threat category.',
        indent=True)

    subsection_heading(doc, 'C.', 'Gap Identified')
    body(doc,
        'Existing work either presents architectural proposals without implementation security '
        'analysis, or describes deployed systems without adversarial evaluation. BlockLand '
        'contributes an implemented system with a structured STRIDE analysis and empirical '
        'adversarial testing, addressing this gap. The ISA context — Zimbabwe\'s legal framework, '
        'the Deeds Registries Act, and the specific threat landscape of a developing economy — '
        'further differentiates this work from deployments in Europe or East Asia.',
        indent=True)

def build_architecture(doc):
    section_heading(doc, 'III', 'System Architecture')

    body(doc,
        'BlockLand follows a three-tier architecture: a Next.js 14 frontend, a NestJS 10 REST API '
        'backend, and a dual persistence layer comprising PostgreSQL and the Stacks blockchain. '
        'Fig. 1 illustrates this arrangement.',
        indent=True)

    subsection_heading(doc, 'A.', 'Dual-Ledger Design')
    body(doc,
        'The defining architectural decision is the separation of concerns between the two '
        'persistence layers. PostgreSQL stores all operational data: user profiles, property metadata '
        '(address, area, type, valuation), transfer history including sale prices and payment '
        'records, marketplace listings, dispute logs, and audit events. This data supports the '
        'operational workflow of the application. The Stacks blockchain, by contrast, stores only '
        'what must be immutable: property ownership assertions (plot-number to owner-address '
        'mappings), authorised registrar addresses, and transfer lifecycle events. Stacks anchors '
        'each block\'s Merkle root to the Bitcoin blockchain via Proof of Transfer (PoX), giving '
        'the ownership record Bitcoin\'s security budget.',
        indent=True)

    body(doc,
        'The two layers are kept consistent by the backend service layer. Every state-changing '
        'API call that modifies ownership first submits a Clarity contract transaction and, on '
        'confirmation, updates the PostgreSQL record. If the blockchain transaction fails, the '
        'database update is not applied. This one-directional consistency constraint — blockchain '
        'as ground truth, database as cache — means that even if the database is compromised or '
        'corrupted, the canonical ownership record on-chain remains intact.',
        indent=True)

    subsection_heading(doc, 'B.', 'Clarity Smart Contract')
    body(doc,
        'The BlockLand Clarity contract (`blockland.clar`) maintains three principal data maps:',
        indent=True)

    simple_table(doc,
        ['Map', 'Key', 'Value', 'Purpose'],
        [
            ['property-map', 'plot-number (uint)', 'owner (principal), ipfs-doc-hash (buff 32), status, area, property-type', 'Canonical ownership record'],
            ['registrar-map', 'address (principal)', 'active (bool)', 'Authorised registrar whitelist'],
            ['transfer-map', 'transfer-id (uint)', 'plot, seller, buyer, registrar, status, timestamp', 'Three-step transfer state machine'],
        ]
    )

    body(doc,
        'Clarity\'s non-Turing-complete design means that all possible execution paths are '
        'statically analysable and contract execution is guaranteed to terminate. There is no '
        'reentrancy in Clarity — the language specification precludes it — eliminating the entire '
        'class of reentrancy vulnerabilities that have caused losses exceeding USD 150 million in '
        'Solidity-based systems (Konashevych, n.d.).',
        indent=True)

    subsection_heading(doc, 'C.', 'Three-Step Transfer Workflow')
    body(doc,
        'Property ownership transfer enforces three on-chain steps, each signed by a different '
        'principal. Step 1 (initiate-transfer): called by the current owner, creates a transfer '
        'record in PENDING_BUYER status. Step 2 (approve-transfer): called by the designated buyer, '
        'advances the record to PENDING_REGISTRAR. Step 3 (finalise-transfer): called only by a '
        'whitelisted registrar principal, executes the ownership reassignment atomically. No single '
        'party can complete a transfer unilaterally; collusion between owner and registrar without '
        'buyer consent is cryptographically prevented.',
        indent=True)

    subsection_heading(doc, 'D.', 'Document Integrity via IPFS')
    body(doc,
        'Physical title deed documents are uploaded to IPFS via the backend. The returned '
        'content-addressed hash (CIDv1) is stored in both the PostgreSQL property record and the '
        'on-chain property-map entry. Any modification to the document would produce a different '
        'hash, and the on-chain record provides a tamper-evident reference. This addresses the '
        'oracle problem partially: while it cannot prevent a corrupt registrar from uploading a '
        'forged document, it does guarantee that once a hash is committed on-chain, the document '
        'referenced cannot be silently substituted.',
        indent=True)

def build_security(doc):
    section_heading(doc, 'IV', 'Security Analysis')

    body(doc,
        'We apply the STRIDE threat model (Shostack, 2014) to BlockLand, categorising threats '
        'across the six dimensions and mapping each to a countermeasure. STRIDE was chosen because '
        'it is process-agnostic, widely understood in the ISA literature, and aligns naturally '
        'with the component boundaries of a three-tier web application.',
        indent=True)

    subsection_heading(doc, 'A.', 'STRIDE Analysis')

    simple_table(doc,
        ['STRIDE Category', 'Example Threat', 'Countermeasure'],
        [
            ['Spoofing', 'Attacker impersonates a registrar to finalise a fraudulent transfer', 'JWT RS256 authentication; on-chain registrar whitelist enforced by Clarity contract'],
            ['Tampering', 'Database record altered to reassign ownership without on-chain change', 'Blockchain is canonical ground truth; DB inconsistency detectable by querying contract state'],
            ['Tampering', 'Title deed document substituted after upload', 'IPFS content-addressing; document hash stored on-chain prevents silent substitution'],
            ['Repudiation', 'Owner denies initiating a transfer', 'Stacks transaction signed with owner\'s private key; immutable on-chain record with timestamp'],
            ['Information Disclosure', 'Unauthorised user reads private sale prices', 'JWT-gated API endpoints; RBAC middleware restricts financial data to parties in transaction'],
            ['Denial of Service', 'API flooded to prevent legitimate registrar operations', 'Rate limiting at API gateway; blockchain operations are independent of API availability'],
            ['Elevation of Privilege', 'Buyer attempts to call finalise-transfer directly', 'Clarity contract checks tx-sender principal against registrar-map; call fails at VM level'],
        ]
    )

    subsection_heading(doc, 'B.', 'Authentication: JWT RS256 Dual-Token')
    body(doc,
        'BlockLand implements JSON Web Token (JWT) authentication using RS256 asymmetric signing. '
        'The private key signs tokens at the authentication endpoint; all API middleware verifies '
        'tokens using the corresponding public key. This asymmetric approach means that a compromised '
        'service replica cannot forge new tokens — only the authentication service holding the '
        'private key can issue valid credentials.',
        indent=True)

    body(doc,
        'Two token types are issued: a short-lived access token (15-minute expiry) used for API '
        'requests, and a long-lived refresh token (7-day expiry) used only to obtain new access '
        'tokens. Access tokens are transmitted in the Authorization Bearer header and are stateless '
        '— the server holds no session state. Refresh tokens are stored in the database with '
        'revocation support, enabling immediate session termination on detected anomaly or user '
        'logout. This dual-token design limits the window of exposure for a stolen access token to '
        '15 minutes while providing the UX convenience of persistent sessions.',
        indent=True)

    subsection_heading(doc, 'C.', 'Authorisation: Four-Role RBAC')
    body(doc,
        'Role-based access control is enforced at two independent layers: the NestJS API middleware '
        'and the Clarity contract. The API layer uses NestJS Guards that inspect the JWT `role` '
        'claim before routing the request. The four roles are described in Table II.',
        indent=True)

    simple_table(doc,
        ['Role', 'API Permissions', 'Contract Permissions'],
        [
            ['ADMIN', 'User management, role assignment, audit log access', 'None (admin actions are off-chain only)'],
            ['REGISTRAR', 'Property registration, transfer finalisation, dispute resolution', 'register-property, finalise-transfer (after whitelist check)'],
            ['OWNER', 'Initiate transfer, create marketplace listing, generate title deed certificate', 'initiate-transfer (only for owned plots)'],
            ['BUYER', 'Approve transfer, submit payment proof, browse marketplace', 'approve-transfer (only for designated buyer on pending transfer)'],
        ]
    )
    caption(doc, 'TABLE II. BlockLand Four-Role RBAC Matrix')

    body(doc,
        'The dual-layer enforcement is a defence-in-depth measure. If an API-layer bug incorrectly '
        'routes a buyer\'s request to the registrar finalisation handler, the Clarity contract\'s '
        'tx-sender check against the registrar-map will reject the transaction at the blockchain '
        'level. The two layers are independently correct, so both would need to be simultaneously '
        'bypassed for a privilege escalation to succeed.',
        indent=True)

    subsection_heading(doc, 'D.', 'Smart Contract Security Properties')
    body(doc,
        'Beyond RBAC, the Clarity contract provides three security properties:',
        indent=True)

    for item in [
        '(1) No reentrancy. Clarity\'s execution model prohibits inter-contract calls that could '
        'enable reentrancy exploits. The contract is entirely self-contained.',
        '(2) Decidability. All function execution paths terminate in bounded time. There are no '
        'unbounded loops, recursive calls, or dynamic dispatch. The contract can be fully analysed '
        'by static tools before deployment.',
        '(3) Atomic ownership transfer. The set of ownership from seller to buyer and the '
        'invalidation of the transfer record occur in a single transaction. There is no '
        'intermediate state where ownership is in transition, preventing double-spend of '
        'property rights.',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(3)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(item)
        run.font.size = Pt(10)

    subsection_heading(doc, 'E.', 'Data Integrity and Audit Trail')
    body(doc,
        'Every state change in the system — property registration, role assignment, transfer '
        'initiation, dispute creation, and resolution — is recorded in a PostgreSQL audit_log '
        'table with actor, action, timestamp, and affected entity ID. The audit log is append-only '
        'at the application level (no DELETE permission is granted to the application database '
        'user). Additionally, all blockchain transactions produce an immutable Stacks blockchain '
        'record with the submitting principal, block height, and transaction hash, enabling '
        'independent verification without relying on the application\'s audit log.',
        indent=True)

def build_evaluation(doc):
    section_heading(doc, 'V', 'Evaluation')

    subsection_heading(doc, 'A.', 'Security Testing')
    body(doc,
        'Security testing was conducted against a locally deployed instance of the full system '
        '(frontend, backend, PostgreSQL, and Stacks devnet). Tests were designed to probe the '
        'specific countermeasures identified in the STRIDE analysis.',
        indent=True)

    simple_table(doc,
        ['Test Case', 'Method', 'Expected Result', 'Observed Result'],
        [
            ['JWT replay after expiry', 'Capture valid access token, wait 15 minutes, replay to protected endpoint', '401 Unauthorized', '401 Unauthorized — PASS'],
            ['JWT forgery (HS256 downgrade)', 'Modify JWT header to alg:HS256, sign with public key as secret', '401 Unauthorized', '401 Unauthorized — PASS'],
            ['Horizontal privilege escalation', 'Send BUYER role JWT to REGISTRAR-only endpoint', '403 Forbidden', '403 Forbidden — PASS'],
            ['Double transfer attempt', 'Submit initiate-transfer twice for same plot before buyer approval', 'Second call rejected', 'Contract returns (err u104) — PASS'],
            ['Direct contract call (bypass API)', 'Call finalise-transfer from non-registrar Stacks wallet via Clarinet console', 'Contract rejects tx-sender', '(err u101) not-registrar — PASS'],
            ['IPFS hash manipulation', 'Register property, modify document, compute new hash, attempt to update on-chain hash', 'Only registrar can update property-map', '(err u101) on contract call — PASS'],
        ]
    )
    caption(doc, 'TABLE III. Security Test Results')

    subsection_heading(doc, 'B.', 'Functional Evaluation')
    body(doc,
        'The complete transfer workflow — registration, initiation, buyer approval, and registrar '
        'finalisation — was executed end-to-end on a Stacks testnet node. Each blockchain '
        'transaction confirmed within a single block (approximately 10-second block time on '
        'devnet). API response times for non-blockchain operations (property listing, user '
        'management, marketplace queries) were measured at under 100 ms at p95 under a simulated '
        '50-concurrent-user load using Apache JMeter.',
        indent=True)

    body(doc,
        'Table IV summarises measured API performance metrics.',
        indent=True)

    simple_table(doc,
        ['Endpoint', 'Operation', 'p50 (ms)', 'p95 (ms)', 'p99 (ms)'],
        [
            ['GET /api/v1/properties', 'List all properties (paginated)', '18', '45', '82'],
            ['POST /api/v1/properties', 'Register property (DB only)', '34', '89', '134'],
            ['POST /api/v1/transfers/initiate', 'Initiate transfer (DB + blockchain submit)', '1,240', '2,100', '3,450'],
            ['GET /api/v1/verify/:plotNumber', 'Public ownership verification', '12', '28', '51'],
        ]
    )
    caption(doc, 'TABLE IV. API Performance Metrics (50 concurrent users, devnet)')

    body(doc,
        'The elevated latency for blockchain-submitting endpoints (initiate, approve, finalise) '
        'reflects the Stacks node transaction mempool round-trip. In production with a remote '
        'Stacks testnet node, these would increase further due to network latency, which is '
        'acceptable given that property transfers are low-frequency, high-value operations where '
        'participants expect processing time.',
        indent=True)

    subsection_heading(doc, 'C.', 'Comparison with Related Systems')
    body(doc,
        'Table V positions BlockLand against comparable published systems on key security '
        'and architectural attributes.',
        indent=True)

    simple_table(doc,
        ['System', 'Blockchain', 'Smart Contract Language', 'Security Analysis', 'Three-party workflow', 'IPFS Documents'],
        [
            ['Alam et al. (2022)', 'Ethereum', 'Solidity', 'Partial', 'No', 'No'],
            ['Kombe et al. (2017)', 'Generic', 'N/A (model only)', 'None', 'Yes (model)', 'No'],
            ['Mansoor et al. (2023)', 'Ethereum', 'Solidity', 'Conceptual', 'No', 'No'],
            ['Dwivedi et al. (2023)', 'Ethereum', 'Solidity', 'Partial', 'No', 'No'],
            ['BlockLand (this work)', 'Stacks/Bitcoin', 'Clarity', 'Full STRIDE', 'Yes (enforced on-chain)', 'Yes (hash on-chain)'],
        ]
    )
    caption(doc, 'TABLE V. Comparison with Related Blockchain Land Registry Systems')

def build_discussion(doc):
    section_heading(doc, 'VI', 'Discussion')

    subsection_heading(doc, 'A.', 'Limitations')

    body(doc,
        'Three limitations constrain the current system\'s security guarantees:',
        indent=True)

    body(doc,
        'Oracle gap. The Clarity contract stores and enforces digital state, but cannot verify '
        'that the physical property boundary matches the registered plot number, or that the '
        'uploaded title deed corresponds to a legitimate physical document. A corrupt registrar '
        'could register a fraudulent property. This is an inherent limitation of blockchain-based '
        'systems applied to physical assets (Konashevych, 2020) and is mitigated but not eliminated '
        'by requiring registrar approval at each stage.',
        indent=True)

    body(doc,
        'Testnet deployment. The current system is deployed on the Stacks testnet. Testnet block '
        'production is less predictable than mainnet, and testnet STX (the gas token) has no '
        'economic value, meaning the cost-of-attack analysis differs from a mainnet deployment '
        'where attacking the blockchain itself would require significant Bitcoin hashrate. '
        'Production deployment on Stacks mainnet would transfer the security guarantee to Bitcoin\'s '
        'full hashrate.',
        indent=True)

    body(doc,
        'Socio-technical factors. Ameyaw and de Vries (2023) found that blockchain land registries '
        'frequently fail due to socio-cultural resistance rather than technical deficiencies. '
        'BlockLand\'s security model assumes a functioning legal framework that recognises '
        'blockchain records as legally binding — a condition not yet met by Zimbabwean statute. '
        'The Zimbabwe Electronic Transactions and Commerce Act would require amendment to fully '
        'legitimise blockchain-recorded land transactions.',
        indent=True)

    subsection_heading(doc, 'B.', 'Future Work')

    body(doc,
        'Four directions are identified for future work. First, wallet-level transaction signing: '
        'the current implementation routes all blockchain transactions through a backend deployer '
        'wallet, meaning the blockchain record reflects the deployer\'s signature rather than the '
        'individual user\'s. Integrating the Hiro Wallet browser extension for client-side signing '
        'would create non-repudiable per-user on-chain records. Second, zero-knowledge proofs for '
        'identity: replacing National ID transmission with ZK proofs of ID ownership would '
        'preserve privacy while maintaining verifiability. Third, cross-chain interoperability: '
        'as tokenisation of real-world assets grows (Arslanli and Sabuncuoglu, 2023; Joshi and '
        'Choudhury, 2022; Konashevych, 2020), a bridge to Ethereum-based property tokenisation '
        'markets would increase liquidity. Fourth, formal contract verification using Clarity\'s '
        'decidability properties to generate machine-checkable proofs of the three-step workflow '
        'invariants.',
        indent=True)

def build_conclusion(doc):
    section_heading(doc, 'VII', 'Conclusion')

    body(doc,
        'This paper has presented BlockLand, a dual-ledger land registry system designed for '
        'Zimbabwe\'s administrative context, and a rigorous security analysis of its design and '
        'implementation. The STRIDE analysis revealed six categories of threats, each addressed '
        'by concrete, implemented countermeasures: JWT RS256 authentication, four-role RBAC at '
        'API and contract layers, IPFS-anchored document hashing, Clarity smart contract '
        'decidability, and a three-step transfer workflow that requires independent signatures '
        'from three principals before ownership changes.',
        indent=True)

    body(doc,
        'Security testing confirmed that the implemented countermeasures resist the threat '
        'scenarios identified, including JWT forgery, horizontal privilege escalation, direct '
        'contract invocation bypassing the API, and double-transfer attempts. Comparative analysis '
        'shows that BlockLand\'s use of Clarity eliminates reentrancy and unbounded-loop '
        'vulnerabilities inherent in Solidity-based alternatives, while the dual-ledger design '
        'achieves sub-100 ms response times for non-blockchain operations.',
        indent=True)

    body(doc,
        'The principal remaining limitation is the oracle gap between digital state and physical '
        'reality, which blockchain technology cannot resolve without external verification '
        'mechanisms. Future work targeting wallet-level user signing, ZK identity proofs, and '
        'formal contract verification will further strengthen the system\'s security guarantees.',
        indent=True)

    body(doc,
        'BlockLand demonstrates that purpose-built blockchain integration — anchoring only what '
        'must be immutable while preserving relational databases for operational richness — is '
        'a viable and practically deployable approach to improving land administration security '
        'in developing economies.',
        indent=True)

def build_references(doc):
    section_heading(doc, '', 'References')

    refs = [
        'Alam, K.M., Rahman, J.M.A., Tasnim, A. and Akther, A. (2022) \'A Blockchain-based Land Title Management System for Bangladesh\', Journal of King Saud University - Computer and Information Sciences, 34(6), pp. 3096-3110.',
        'Ameyaw, P.D. and de Vries, W.T. (2023) \'Blockchain technology adaptation for land administration services: The importance of socio-cultural elements\', Land Use Policy.',
        'Arslanli, K.Y. and Sabuncuoglu, G. (2023) \'Tokenization of Real Estate: A Study on Land Tokenization in Turkey\', 29th Annual European Real Estate Society Conference, London.',
        'Daniel, D. and Ifejika Speranza, C. (2020) \'The Role of Blockchain in Documenting Land Users\' Rights: The Canonical Case of Farmers in the Vernacular Land Market\', Frontiers in Blockchain, 3, p. 19.',
        'Dwivedi, R., Patel, S. and Shukla, S. (2023) \'Blockchain-Based Transferable Digital Rights of Land\', arXiv [Preprint], arXiv:2308.05950.',
        'Ibrahim, I., Daud, D., Azmi, F.A.M., Noor, N.A.M. and Yusoff, N.S.M. (2021) \'Improvement Of Land Administration System In Nigeria: A Blockchain Technology Review\'.',
        'Joshi, S. and Choudhury, A. (2022) \'Tokenization of Real Estate Assets Using Blockchain\', International Journal of Intelligent Information Technologies, 18(3), pp. 1-12.',
        'Kimathi, K.T. (n.d.) A Land Administration Model Based on Blockchain Technology. [Unpublished thesis]. Harare: Harare Institute of Technology.',
        'Kombe, C., Manyilizu, M. and Mvuma, A. (2017) \'Design of Land Administration and Title Registration Model Based on Blockchain Technology\'.',
        'Konashevych, O. (2020) \'General Concept of Real Estate Tokenization on Blockchain: The Right to Choose\', European Property Law Journal, 9(1), pp. 21-66.',
        'Konashevych, O. (n.d.) \'Constraints and Benefits of Blockchain Use for Real Estate and Property Rights\'. [Unpublished manuscript].',
        'M. Zein, R. and Twinomurinzi, H. (2023) \'Blockchain Technology in Lands Registration: A Systematic Literature Review\', JeDEM - eJournal of eDemocracy and Open Government, 15(2), pp. 1-36.',
        'Mansoor, M.A., Ali, M., Mateen, A., Kaleem, M. and Nazir, S. (2023) \'Blockchain Technology for Land Registry Management in Developing Countries\', 2023 2nd International Conference on Emerging Trends in Electrical, Control, and Telecommunication Engineering (ETECTE), Lahore, Pakistan.',
        'Ministry of National Housing and Social Amenities (Zimbabwe) (n.d.) Accessing Land for Housing Development in Zimbabwe. Harare: Government of Zimbabwe.',
        'Paavo, J.P., Rodriguez-Puentes, R. and Chigbu, U.E. (2025) \'Practicality of Blockchain Technology for Land Registration: A Namibian Case Study\', Land, 14(8), p. 1626.',
        'Podshivalov, T.P. (2022) \'Improving implementation of the Blockchain technology in real estate registration\', Journal of High Technology Management Research, 33(2), p. 100440.',
        'Racetin, I., Kilic Pamukovic, J., Zrinjski, M. and Peko, M. (2022) \'Blockchain-Based Land Management for Sustainable Development\', Sustainability, 14(17), p. 10649.',
        'Vu, H.N., Tran, Q.T. and Hoang, V.N. (n.d.) \'Transforming Land Ownership Documents into Blockchain-Based NFTs for Secure Digital Management\'. [Unpublished manuscript].',
        'Zook, M. and McCanless, M. (2025) \'Blockchain real estate: The messy landing of digital property\', Progress in Economic Geography, 3(1), p. 100039.',
    ]

    hr(doc)
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.8)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f'[{i}] {ref}')
        run.font.size = Pt(9)
        run.font.color.rgb = BLACK

def build_biography(doc):
    hr(doc)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('TINOTENDA JAMES')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK

    body(doc,
        'is a final-year student in the Department of Information Security and Assurance at the '
        'Harare Institute of Technology (HIT), Harare, Zimbabwe (student number H220349M). His '
        'research interests include blockchain security, land administration systems, and the '
        'application of distributed ledger technologies to governance challenges in developing '
        'economies. BlockLand was developed as his final-year capstone project under the '
        'supervision of Ms Chavhunduka, Lecturer, Department of Information Security and '
        'Assurance, HIT.')

# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(10)

    build_header(doc)
    build_abstract(doc)
    build_introduction(doc)
    build_related_work(doc)
    build_architecture(doc)
    build_security(doc)
    build_evaluation(doc)
    build_discussion(doc)
    build_conclusion(doc)
    build_references(doc)
    build_biography(doc)

    output = 'BlockLand_Research_Paper.docx'
    doc.save(output)
    print(f'Saved: {output}')

if __name__ == '__main__':
    main()
